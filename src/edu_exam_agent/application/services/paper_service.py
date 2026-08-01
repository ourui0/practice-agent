"""Question selection and DOCX export for exams and practice sets."""

from __future__ import annotations

import json
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from edu_exam_agent.application.services.document_service import DocumentService
from edu_exam_agent.application.services.question_bank_service import QuestionBankService
from edu_exam_agent.application.services.question_types import (
    QUESTION_TYPE_LABELS,
    QUESTION_TYPE_ORDER,
    ordered_type_counts,
)
from edu_exam_agent.infrastructure.database.models import (
    PaperHistoryItemModel,
    PaperHistoryModel,
    QuestionModel,
)


@dataclass(frozen=True, slots=True)
class PaperRequest:
    course_id: int
    title: str
    question_types: tuple[str, ...]
    count: int
    target_difficulty: int | None = None
    minimum_score: float = 0
    include_answers: bool = True
    duration_minutes: int = 90
    document_id: int | None = None
    chapter_ids: tuple[int, ...] = ()
    question_type_counts: tuple[tuple[str, int], ...] = ()
    exclude_recent_days: int = 180
    exclude_recent_papers: int = 20

    def validate(self) -> None:
        if not self.title.strip():
            raise ValueError("标题不能为空")
        if self.count < 1 or self.count > 200:
            raise ValueError("题目数量必须在 1 到 200 之间")
        if self.question_type_counts:
            configured_types = [item[0] for item in self.question_type_counts]
            if len(configured_types) != len(set(configured_types)):
                raise ValueError("题型数量配置中存在重复题型")
            unknown = [
                question_type
                for question_type in configured_types
                if question_type not in QUESTION_TYPE_ORDER
            ]
            if unknown:
                raise ValueError(f"不支持的题型：{'、'.join(unknown)}")
            if any(count < 0 for _, count in self.question_type_counts):
                raise ValueError("题型数量不能为负数")
            configured_total = sum(count for _, count in self.question_type_counts)
            if configured_total < 1:
                raise ValueError("题目总数必须大于0")
            if configured_total != self.count:
                raise ValueError("各题型数量之和与题目总数不一致")
        elif not self.question_types:
            raise ValueError("请至少选择一种题型")
        if self.target_difficulty is not None and self.target_difficulty not in range(1, 6):
            raise ValueError("目标难度必须在 1 到 5 之间")
        if self.exclude_recent_days < 0 or self.exclude_recent_papers < 0:
            raise ValueError("近期试卷排除范围不能为负数")


@dataclass(frozen=True, slots=True)
class Paper:
    title: str
    questions: tuple[QuestionModel, ...]
    duration_minutes: int
    include_answers: bool
    history_id: int | None = None

    @property
    def total_score(self) -> int:
        return sum(question.score for question in self.questions)


class PaperService:
    def __init__(self, bank: QuestionBankService) -> None:
        self._bank = bank

    def assemble(self, request: PaperRequest) -> Paper:
        request.validate()
        candidates = self._candidates(request)
        if request.question_type_counts:
            selected = self._select_by_quota(request, candidates)
        elif len(candidates) < request.count:
            raise ValueError(
                f"符合条件且教材边界通过的题目只有 {len(candidates)} 道，"
                f"不足以生成 {request.count} 道题"
            )
        else:
            selected = candidates[: request.count]
            selected.sort(key=self._type_order_key)
        history_id = self._record_draft(request, selected)
        return Paper(
            request.title.strip(),
            tuple(selected),
            request.duration_minutes,
            request.include_answers,
            history_id,
        )

    def available_count(self, request: PaperRequest) -> int:
        request.validate()
        return len(self._candidates(request))

    def available_count_by_type(self, request: PaperRequest) -> dict[str, int]:
        request.validate()
        counts = {question_type: 0 for question_type in QUESTION_TYPE_ORDER}
        for question in self._candidates(request):
            if question.question_type in counts:
                counts[question.question_type] += 1
        return counts

    def _candidates(self, request: PaperRequest) -> list[QuestionModel]:
        if request.document_id is not None:
            DocumentService(self._bank.engine).assert_ready_for_generation(
                request.document_id
            )
        effective_types = (
            tuple(question_type for question_type, _ in ordered_type_counts(
                request.question_type_counts
            ))
            if request.question_type_counts
            else request.question_types
        )
        recent_ids = self._recent_question_ids(request)
        candidates = [
            question
            for question in self._bank.list(
                course_id=request.course_id,
                minimum_score=request.minimum_score,
                document_id=request.document_id,
                chapter_ids=request.chapter_ids,
            )
            if question.question_type in effective_types
            and question.id not in recent_ids
            and question.boundary_passed
            and question.status in {"validated", "teacher_edited"}
        ]
        candidates.sort(key=lambda question: self._candidate_rank(request, question))
        return candidates

    def _recent_question_ids(self, request: PaperRequest) -> set[int]:
        if request.exclude_recent_days == 0 or request.exclude_recent_papers == 0:
            return set()
        cutoff = datetime.now() - timedelta(days=request.exclude_recent_days)
        with Session(self._bank.engine) as session:
            paper_ids = list(
                session.scalars(
                    select(PaperHistoryModel.id)
                    .where(
                        PaperHistoryModel.course_id == request.course_id,
                        PaperHistoryModel.status.in_(("exported", "used")),
                        PaperHistoryModel.created_at >= cutoff,
                    )
                    .order_by(PaperHistoryModel.id.desc())
                    .limit(request.exclude_recent_papers)
                )
            )
            if not paper_ids:
                return set()
            return set(
                session.scalars(
                    select(PaperHistoryItemModel.question_id).where(
                        PaperHistoryItemModel.paper_id.in_(paper_ids)
                    )
                )
            )

    def _record_draft(
        self, request: PaperRequest, selected: list[QuestionModel]
    ) -> int:
        with Session(self._bank.engine) as session, session.begin():
            history = PaperHistoryModel(
                course_id=request.course_id,
                title=request.title.strip(),
                status="draft",
                request_json=json.dumps(asdict(request), ensure_ascii=False),
            )
            session.add(history)
            session.flush()
            for position, question in enumerate(selected, 1):
                session.add(
                    PaperHistoryItemModel(
                        paper_id=history.id,
                        question_id=question.id,
                        position=position,
                        snapshot_json=json.dumps(
                            self._bank._snapshot(question), ensure_ascii=False
                        ),
                    )
                )
            return history.id

    def mark_used(self, history_id: int) -> None:
        with Session(self._bank.engine) as session, session.begin():
            history = session.get(PaperHistoryModel, history_id)
            if history is None:
                raise ValueError("试卷历史不存在")
            history.status = "used"
            history.used_at = datetime.now()

    def load(self, history_id: int) -> Paper:
        """Restore a persisted paper so chat and assembly pages can preview it."""
        with Session(self._bank.engine) as session:
            history = session.get(PaperHistoryModel, history_id)
            if history is None:
                raise ValueError("试卷历史不存在")
            items = list(
                session.scalars(
                    select(PaperHistoryItemModel)
                    .where(PaperHistoryItemModel.paper_id == history_id)
                    .order_by(PaperHistoryItemModel.position)
                )
            )
            questions = [
                session.get(QuestionModel, item.question_id) for item in items
            ]
            if any(question is None for question in questions):
                raise ValueError("试卷中的部分题目已经不存在")
            request = json.loads(history.request_json or "{}")
            for question in questions:
                session.expunge(question)
            return Paper(
                history.title,
                tuple(questions),
                int(request.get("duration_minutes", 90)),
                bool(request.get("include_answers", True)),
                history.id,
            )

    def _select_by_quota(
        self, request: PaperRequest, candidates: list[QuestionModel]
    ) -> list[QuestionModel]:
        requested = dict(ordered_type_counts(request.question_type_counts))
        selected: list[QuestionModel] = []
        for question_type in QUESTION_TYPE_ORDER:
            quota = requested.get(question_type, 0)
            if quota <= 0:
                continue
            type_candidates = [
                question
                for question in candidates
                if question.question_type == question_type
            ]
            if len(type_candidates) < quota:
                missing = quota - len(type_candidates)
                label = QUESTION_TYPE_LABELS[question_type]
                raise ValueError(
                    f"{label}需要{quota}道，但当前只有{len(type_candidates)}道可用，"
                    f"还缺{missing}道"
                )
            selected.extend(type_candidates[:quota])
        return selected

    @staticmethod
    def _candidate_rank(request: PaperRequest, question: QuestionModel) -> tuple:
        difficulty_gap = (
            abs(question.difficulty - request.target_difficulty)
            if request.target_difficulty is not None
            else 0
        )
        return (
            difficulty_gap,
            -(question.recommendation_score or 0),
            -(question.quality_score or 0),
            -question.id,
        )

    @staticmethod
    def _type_order_key(question: QuestionModel) -> tuple[int, int]:
        try:
            order = QUESTION_TYPE_ORDER.index(question.question_type)
        except ValueError:
            order = len(QUESTION_TYPE_ORDER)
        return order, 0

    def preview(self, paper: Paper, include_answers: bool = False) -> str:
        lines = [
            paper.title,
            (
                f"共 {len(paper.questions)} 题　总分 {paper.total_score} 分　"
                f"建议时长 {paper.duration_minutes} 分钟"
            ),
            "",
        ]
        current_type = ""
        section_number = 0
        type_totals = _type_totals(paper.questions)
        for index, question in enumerate(paper.questions, 1):
            if question.question_type != current_type:
                current_type = question.question_type
                section_number += 1
                lines.extend(
                    (
                        _section_title(
                            section_number,
                            current_type,
                            type_totals[current_type],
                        ),
                        "",
                    )
                )
            lines.append(f"{index}. {question.stem}（{question.score}分）")
            if self._bank.figure(question.id) is not None:
                lines.append("   [本题含配图，导出 Word 时自动插入]")
            for option in _options(question):
                lines.append(f"   {option['label']}. {option['content']}")
            if include_answers:
                lines.extend((f"   答案：{question.answer}", f"   解析：{question.analysis}"))
            lines.append("")
        return "\n".join(lines)

    def export_docx(self, paper: Paper, path: Path) -> Path:
        document = Document()
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        _configure_styles(document)
        _add_header_footer(section, paper.title)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(paper.title)
        run.bold = True
        run.font.size = Pt(20)
        _set_run_font(run, "Microsoft YaHei")
        meta = document.add_paragraph(
            f"考试时间：{paper.duration_minutes} 分钟　　满分：{paper.total_score} 分"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            "姓名：________________　班级：________________　得分：________________"
        )

        current_type = ""
        section_number = 0
        type_totals = _type_totals(paper.questions)
        for index, question in enumerate(paper.questions, 1):
            if question.question_type != current_type:
                current_type = question.question_type
                section_number += 1
                document.add_heading(
                    _section_title(
                        section_number,
                        current_type,
                        type_totals[current_type],
                    ),
                    level=1,
                )
            paragraph = document.add_paragraph(style="Question")
            paragraph.add_run(f"{index}. {question.stem}").bold = True
            paragraph.add_run(f"（{question.score}分）")
            for option in _options(question):
                document.add_paragraph(
                    f"{option['label']}. {option['content']}", style="Option"
                )
            figure = self._bank.figure(question.id)
            if figure is not None:
                picture = document.add_picture(BytesIO(figure.png_data), width=Cm(12.5))
                picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if question.question_type not in ("单项选择题", "多项选择题", "判断题"):
                document.add_paragraph("答：\n\n", style="AnswerSpace")

        if paper.include_answers:
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_heading("参考答案与解析", level=1)
            current_type = ""
            section_number = 0
            for index, question in enumerate(paper.questions, 1):
                if question.question_type != current_type:
                    current_type = question.question_type
                    section_number += 1
                    document.add_heading(
                        _section_title(
                            section_number,
                            current_type,
                            type_totals[current_type],
                        ),
                        level=2,
                    )
                answer = document.add_paragraph(style="Question")
                answer.add_run(f"{index}. 答案：").bold = True
                answer.add_run(question.answer)
                analysis = document.add_paragraph(style="Analysis")
                analysis.add_run("解析：").bold = True
                analysis.add_run(question.analysis)
                if question.scoring_criteria:
                    criteria = document.add_paragraph(style="Analysis")
                    criteria.add_run("评分标准：").bold = True
                    criteria.add_run(question.scoring_criteria)

        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
        if paper.history_id is not None:
            with Session(self._bank.engine) as session, session.begin():
                history = session.get(PaperHistoryModel, paper.history_id)
                if history is not None:
                    session.execute(
                        delete(PaperHistoryItemModel).where(
                            PaperHistoryItemModel.paper_id == paper.history_id
                        )
                    )
                    for position, question in enumerate(paper.questions, 1):
                        session.add(
                            PaperHistoryItemModel(
                                paper_id=paper.history_id,
                                question_id=question.id,
                                position=position,
                                snapshot_json=json.dumps(
                                    self._bank._snapshot(question), ensure_ascii=False
                                ),
                            )
                        )
                    history.status = "exported"
                    history.exported_at = datetime.now()
        return path


def _options(question: QuestionModel) -> list[dict[str, str]]:
    try:
        value = json.loads(question.options_json)
    except (TypeError, json.JSONDecodeError):
        return []
    return value if isinstance(value, list) else []


def _type_totals(questions: tuple[QuestionModel, ...]) -> dict[str, int]:
    totals: dict[str, int] = {}
    for question in questions:
        totals[question.question_type] = totals.get(question.question_type, 0) + 1
    return totals


def _section_title(number: int, question_type: str, count: int) -> str:
    numerals = ("一", "二", "三", "四", "五", "六", "七", "八", "九", "十")
    prefix = numerals[number - 1] if number <= len(numerals) else str(number)
    label = QUESTION_TYPE_LABELS.get(question_type, question_type)
    return f"{prefix}、{label}（共{count}题）"


def _set_run_font(run, name: str) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, left, first, after in (
        ("Question", 0, 0, 7),
        ("Option", 18, 0, 3),
        ("AnswerSpace", 18, 0, 6),
        ("Analysis", 18, 0, 5),
    ):
        style = document.styles.add_style(name, 1)
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Pt(left)
        style.paragraph_format.first_line_indent = Pt(first)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    heading = document.styles["Heading 1"]
    heading.font.name = "Microsoft YaHei"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(31, 78, 121)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_header_footer(section, title: str) -> None:
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        _set_run_font(run, "Microsoft YaHei")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    footer.add_run(" 页")
