"""Textbook-grounded learning-guide and teaching-plan generation."""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from sqlalchemy import Engine, select
from sqlalchemy.orm import Session

from edu_exam_agent.application.services.document_service import DocumentService
from edu_exam_agent.infrastructure.database.models import (
    ChapterModel,
    CourseModel,
    DocumentChunkModel,
    DocumentModel,
    KnowledgePointModel,
    TeachingPackageModel,
)
from edu_exam_agent.infrastructure.llm.provider import LLMProvider
from edu_exam_agent.infrastructure.retrieval import FtsRetriever, SearchResult

SYSTEM_PROMPT = """你是一名熟悉中国中小学课程标准、课堂教学设计和教材分析的资深教师。
请根据提供的教材证据和目标知识点，生成相互对应的学生《导学案》和教师《教案》。

严格规则：
1. 只能依据教材证据中的概念、公式、定理、方法、例题思想和学习要求，不得使用模型记忆补写教材事实。
2. 可以原创课堂情境和练习，但不得引入超出学段、章节及指定知识点的新知识。
3. 每个学习目标、任务、练习和教学环节都要关联目标知识点与真实 evidence_id；不得编造证据编号或页码。
4. 导学案与教案使用相同的目标和任务编号。
   导学案中的当堂练习和课后任务必须直接提供答案与简明解析，便于学生自查。
   教案 answer_reference 还要提供完整解析、评分标准、常见错误和纠正策略。
5. 学习目标必须具体、可观察、可评价。教学过程时间之和必须等于课时长度。
6. 证据不足时缩小范围，并在 insufficiencies 中说明；完全不足时返回 insufficient，不得编造正文。
7. 不要大段照抄教材。涉及图片、表格、实验或图形而证据没有具体内容时，不得猜测。
8. 数理内容必须保证计算、符号、单位和逻辑正确；开放性问题不得伪装成唯一答案。
9. 单项选择题必须且只能有一个正确选项；生成后逐题复算并排除多解、无解和条件不足。
10. 只要导学案包含练习或课后任务，对应条目的 answer 和 analysis 就不得为空；
    教案 answer_reference 也不得为空，且必须逐题给出答案、解析和评分标准。
11. 导学案显示文字尽量使用规范中文，公式、变量、单位和任务编号除外；不得把 JSON 字段名显示给学生。
12. 导学案必须安排 6 至 10 个知识梳理填空，并在末段单列“知识点总结”。
13. 导学案正文不显示题目难度和分值。所有课前问题、学习任务问题、知识填空、
    当堂练习和课后任务的答案与解析统一放在导学案最后，不得紧跟题目出现。

只输出一个合法 JSON 对象，不要输出 Markdown 或解释。必须包含：
- status: complete、partial 或 insufficient
- title
- basic_info
- material_tracking: knowledge_points、coverage_matrix、unsupported_knowledge_points
- learning_guide: document_name、learning_objectives、key_points、difficult_points、
  prior_knowledge_check、pre_class_preview、knowledge_framework、guided_fill_ins、
  learning_tasks、in_class_practice、knowledge_summary、learning_summary、after_class_tasks
- teaching_plan: document_name、textbook_analysis、student_analysis、teaching_objectives、
  key_points、difficult_points、teaching_methods、learning_methods、teaching_resources、
  teaching_process、differentiated_instruction、board_design、answer_reference、
  homework_design、post_lesson_reflection_template
- quality_check: all_knowledge_points_supported、guide_and_plan_aligned、
  lesson_time_total_minutes、lesson_time_matches、answers_verified、
  out_of_scope_content_found、issues
- insufficiencies

编号规则：目标 O1...，预备诊断 P1...，学习任务 T1...，评价 A1...，作业 H1...。
material_tracking.coverage_matrix 的每项必须建立“知识点—目标—学生任务—评价—教材证据”的对应关系。
learning_objectives 每项包含 id、content、knowledge_points、success_criteria。
pre_class_preview 每项包含 id、instruction、student_output、knowledge_points、evidence_ids。
guided_fill_ins 必须包含 6 至 10 项，每项包含 id（B1...）、prompt（带______）、
answer、analysis、knowledge_points、evidence_ids。
learning_tasks 每项包含 id、title、task_type、scenario_or_material、questions、
learning_hint、student_output、answer、analysis、knowledge_points、evidence_ids。
in_class_practice 每项包含 id、question_type、question、difficulty、
knowledge_points、evidence_ids、score、answer、analysis。
after_class_tasks 每项包含 id、content、level、knowledge_points、
estimated_minutes、answer、analysis；不得引用输入中没有提供的教材题号或未定义的问题编号。
teaching_process 每项必须包含 stage、duration_minutes、related_task_ids、
teacher_activities、student_activities、expected_student_responses、possible_difficulties、
teacher_responses、assessment_id、assessment_method、design_intention、
knowledge_points、evidence_ids。"""


@dataclass(frozen=True, slots=True)
class TeachingPackageRequest:
    course_id: int
    document_id: int
    chapter_ids: tuple[int, ...]
    knowledge_point_ids: tuple[int, ...]
    lesson_type: str = "新授课"
    lesson_duration_minutes: int = 45
    student_profile: str = ""
    teaching_focus: str = ""
    additional_requirements: str = ""

    def validate(self) -> None:
        if self.course_id < 1:
            raise ValueError("请选择课程")
        if self.document_id < 1:
            raise ValueError("请选择教材")
        if not self.chapter_ids:
            raise ValueError("请选择教材章节")
        if not self.knowledge_point_ids:
            raise ValueError("请至少选择一个已确认知识点")
        if not 20 <= self.lesson_duration_minutes <= 180:
            raise ValueError("课时长度应在 20 到 180 分钟之间")
        if not self.lesson_type.strip():
            raise ValueError("请选择课型")


@dataclass(frozen=True, slots=True)
class TeachingEvidence:
    evidence_id: str
    chunk_id: int
    document_name: str
    chapter_title: str
    page_start: int
    page_end: int
    excerpt: str
    knowledge_points: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class TeachingPackageResult:
    record_id: int
    title: str
    status: str
    payload: dict[str, Any]
    evidence: tuple[TeachingEvidence, ...]
    model_name: str


class TeachingPackageService:
    """Generate, validate, persist, render and export a teaching package."""

    def __init__(
        self,
        engine: Engine,
        retriever: FtsRetriever,
        provider: LLMProvider,
        model_name: str,
    ) -> None:
        self._engine = engine
        self._retriever = retriever
        self._provider = provider
        self._model_name = model_name

    def generate(self, request: TeachingPackageRequest) -> TeachingPackageResult:
        request.validate()
        context = self._load_context(request)
        DocumentService(self._engine).assert_ready_for_generation(request.document_id)
        evidence = self._collect_evidence(request, context["points"])
        if not evidence:
            raise ValueError("所选知识点在当前教材章节中没有找到可用依据")

        prompt = self._build_user_prompt(request, context, evidence)
        try:
            payload = self._provider.generate_json(SYSTEM_PROMPT, prompt)
        except TimeoutError as exc:
            raise ValueError(
                "导学案和教案内容较长，模型在限定时间内没有生成完成。"
                "请稍后重试，或减少同时选择的知识点数量。"
            ) from exc
        payload = self._validate_and_normalize(payload, request, context, evidence)
        record_id = self._save(request, payload, evidence)
        return TeachingPackageResult(
            record_id=record_id,
            title=str(payload["title"]),
            status=str(payload["status"]),
            payload=payload,
            evidence=evidence,
            model_name=self._model_name,
        )

    def list_history(
        self, course_id: int | None = None, limit: int = 50
    ) -> list[TeachingPackageModel]:
        with Session(self._engine) as session:
            statement = select(TeachingPackageModel)
            if course_id is not None:
                statement = statement.where(TeachingPackageModel.course_id == course_id)
            statement = statement.order_by(
                TeachingPackageModel.created_at.desc(), TeachingPackageModel.id.desc()
            ).limit(min(max(limit, 1), 200))
            rows = list(session.scalars(statement))
            for row in rows:
                session.expunge(row)
            return rows

    def load(self, record_id: int) -> TeachingPackageResult:
        with Session(self._engine) as session:
            row = session.get(TeachingPackageModel, record_id)
            if row is None:
                raise ValueError("备课记录不存在")
            payload = json.loads(row.result_json)
            request_payload = json.loads(row.request_json)
            evidence = tuple(
                TeachingEvidence(
                    evidence_id=item["evidence_id"],
                    chunk_id=int(item["chunk_id"]),
                    document_name=item["document_name"],
                    chapter_title=item["chapter_title"],
                    page_start=int(item["page_start"]),
                    page_end=int(item["page_end"]),
                    excerpt=item["excerpt"],
                    knowledge_points=tuple(item.get("knowledge_points", [])),
                )
                for item in request_payload.get("evidence", [])
            )
            return TeachingPackageResult(
                row.id, row.title, row.status, payload, evidence, row.model_name
            )

    def delete(self, record_id: int) -> None:
        with Session(self._engine) as session:
            row = session.get(TeachingPackageModel, record_id)
            if row is not None:
                session.delete(row)
                session.commit()

    def export_docx(self, result: TeachingPackageResult, output: Path) -> Path:
        from docx import Document

        if output.suffix.lower() != ".docx":
            output = output.with_suffix(".docx")
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(result.title, 0)
        basic = result.payload.get("basic_info", {})
        document.add_paragraph(
            " · ".join(
                str(value)
                for value in (
                    basic.get("subject"),
                    basic.get("grade"),
                    basic.get("textbook_name"),
                    basic.get("chapter_title"),
                )
                if value
            )
        )
        self._append_text_document(
            document,
            self.render_learning_guide(result.payload, include_answers=False),
            "导学案",
        )
        document.add_page_break()
        self._append_text_document(document, self.render_teaching_plan(result.payload), "教案")
        document.add_page_break()
        self._append_text_document(
            document, self.render_material_tracking(result.payload, result.evidence), "教材依据"
        )
        document.add_page_break()
        self._append_text_document(
            document,
            self.render_learning_guide_answers(result.payload),
            "导学案答案",
        )
        document.save(output)
        return output

    @staticmethod
    def render_learning_guide(
        payload: dict[str, Any], *, include_answers: bool = True
    ) -> str:
        guide = payload.get("learning_guide") or {}
        lines = [str(guide.get("document_name") or "导学案")]
        hidden = {
            "answer",
            "answers",
            "analysis",
            "difficulty",
            "score",
            "question_type",
        }
        TeachingPackageService._append_objectives(
            lines, "学习目标", guide.get("learning_objectives")
        )
        TeachingPackageService._append_simple_list(lines, "学习重点", guide.get("key_points"))
        TeachingPackageService._append_simple_list(lines, "学习难点", guide.get("difficult_points"))
        TeachingPackageService._append_cards(
            lines, "课前诊断", guide.get("prior_knowledge_check"), hidden
        )
        TeachingPackageService._append_cards(
            lines, "课前预习", guide.get("pre_class_preview"), hidden
        )
        TeachingPackageService._append_cards(
            lines, "知识框架", guide.get("knowledge_framework"), hidden
        )
        TeachingPackageService._append_cards(
            lines, "知识梳理填空", guide.get("guided_fill_ins"), hidden
        )
        TeachingPackageService._append_cards(
            lines, "学习任务", guide.get("learning_tasks"), hidden
        )
        TeachingPackageService._append_cards(
            lines, "当堂练习", guide.get("in_class_practice"), hidden
        )
        TeachingPackageService._append_simple_list(
            lines, "知识点总结", guide.get("knowledge_summary")
        )
        summary = guide.get("learning_summary") or {}
        TeachingPackageService._append_simple_list(
            lines, "知识结构", summary.get("knowledge_structure")
        )
        TeachingPackageService._append_simple_list(
            lines, "学习反思", summary.get("student_reflection_prompts")
        )
        TeachingPackageService._append_cards(
            lines, "课后任务", guide.get("after_class_tasks"), hidden
        )
        if include_answers:
            answer_text = TeachingPackageService.render_learning_guide_answers(payload)
            if answer_text:
                lines.extend(["", answer_text])
        return "\n".join(lines).strip()

    @staticmethod
    def render_learning_guide_answers(payload: dict[str, Any]) -> str:
        guide = payload.get("learning_guide") or {}
        lines = ["答案与解析"]
        sections = (
            ("课前诊断答案", "prior_knowledge_check"),
            ("课前预习答案", "pre_class_preview"),
            ("知识框架答案", "knowledge_framework"),
            ("知识梳理填空答案", "guided_fill_ins"),
            ("学习任务答案", "learning_tasks"),
            ("当堂练习答案", "in_class_practice"),
            ("课后任务答案", "after_class_tasks"),
        )
        for title, key in sections:
            items = [
                item
                for item in guide.get(key, [])
                if isinstance(item, dict) and (item.get("answer") or item.get("answers"))
            ]
            if not items:
                continue
            lines.extend(["", title])
            for item in items:
                identifier = item.get("id") or item.get("title") or "•"
                answer = item.get("answer") or item.get("answers")
                lines.append(f"{identifier}  答案：{TeachingPackageService._display(answer)}")
                analysis = item.get("analysis")
                if analysis:
                    lines.append(f"  简析：{TeachingPackageService._display(analysis)}")
        return "\n".join(lines).strip() if len(lines) > 1 else ""

    @staticmethod
    def render_teaching_plan(payload: dict[str, Any]) -> str:
        plan = payload.get("teaching_plan") or {}
        lines = [str(plan.get("document_name") or "教案")]
        TeachingPackageService._append_paragraph(lines, "教材分析", plan.get("textbook_analysis"))
        TeachingPackageService._append_paragraph(lines, "学情分析", plan.get("student_analysis"))
        TeachingPackageService._append_objectives(
            lines, "教学目标", plan.get("teaching_objectives")
        )
        TeachingPackageService._append_cards(lines, "教学重点", plan.get("key_points"))
        TeachingPackageService._append_cards(lines, "教学难点", plan.get("difficult_points"))
        TeachingPackageService._append_simple_list(lines, "教学方法", plan.get("teaching_methods"))
        TeachingPackageService._append_simple_list(lines, "学习方法", plan.get("learning_methods"))
        TeachingPackageService._append_simple_list(
            lines, "教学资源", plan.get("teaching_resources")
        )
        TeachingPackageService._append_cards(lines, "教学过程", plan.get("teaching_process"))
        TeachingPackageService._append_cards(
            lines, "分层教学", [plan.get("differentiated_instruction") or {}]
        )
        TeachingPackageService._append_simple_list(lines, "板书设计", plan.get("board_design"))
        TeachingPackageService._append_cards(lines, "答案与评价参考", plan.get("answer_reference"))
        TeachingPackageService._append_cards(lines, "作业设计", plan.get("homework_design"))
        TeachingPackageService._append_simple_list(
            lines, "课后反思", plan.get("post_lesson_reflection_template")
        )
        return "\n".join(lines).strip()

    @staticmethod
    def render_material_tracking(
        payload: dict[str, Any], evidence: tuple[TeachingEvidence, ...]
    ) -> str:
        tracking = payload.get("material_tracking") or {}
        lines = ["教材追踪"]
        TeachingPackageService._append_tracking_knowledge_points(
            lines, tracking.get("knowledge_points")
        )
        TeachingPackageService._append_tracking_coverage(
            lines, tracking.get("coverage_matrix")
        )
        unsupported = tracking.get("unsupported_knowledge_points") or []
        if unsupported:
            lines.extend(["", "证据不足知识点"])
            for index, item in enumerate(unsupported, start=1):
                lines.append(f"知识点{index}  {TeachingPackageService._display(item)}")
        lines.extend(["", "教材原文依据"])
        for index, item in enumerate(evidence):
            if index:
                lines.append("")
            pages = (
                f"第{item.page_start}页"
                if item.page_start == item.page_end
                else f"第{item.page_start}-{item.page_end}页"
            )
            points = "、".join(item.knowledge_points)
            lines.append(
                f"{item.evidence_id}  《{item.document_name}》 "
                f"{item.chapter_title}（{pages}）"
            )
            if points:
                lines.append(f"  对应知识点：{points}")
            lines.append(f"  教材原文：{item.excerpt}")
        insufficiencies = payload.get("insufficiencies") or []
        TeachingPackageService._append_tracking_insufficiencies(lines, insufficiencies)
        return "\n".join(lines).strip()

    @staticmethod
    def _append_tracking_knowledge_points(lines: list[str], values: Any) -> None:
        if not values:
            return
        lines.extend(["", "知识点覆盖"])
        for index, item in enumerate(values, start=1):
            if index > 1:
                lines.append("")
            if not isinstance(item, dict):
                lines.append(f"知识点{index}  {TeachingPackageService._display(item)}")
                continue
            name = TeachingPackageService._display(item.get("knowledge_point") or "未命名")
            lines.append(f"知识点{index}  {name}")
            TeachingPackageService._append_tracking_field(lines, "作用", item.get("role"))
            TeachingPackageService._append_tracking_field(
                lines, "教学层级", item.get("teaching_level")
            )
            TeachingPackageService._append_tracking_field(
                lines, "教材概述", item.get("source_summary")
            )
            TeachingPackageService._append_tracking_field(
                lines, "教材页码", item.get("page_references")
            )
            TeachingPackageService._append_tracking_field(
                lines, "教材证据", item.get("evidence_ids")
            )

    @staticmethod
    def _append_tracking_coverage(lines: list[str], values: Any) -> None:
        if not values:
            return
        lines.extend(["", "目标与任务对应"])
        for index, item in enumerate(values, start=1):
            if index > 1:
                lines.append("")
            if not isinstance(item, dict):
                lines.append(f"对应关系{index}  {TeachingPackageService._display(item)}")
                continue
            name = TeachingPackageService._display(item.get("knowledge_point") or "未命名")
            lines.append(f"对应关系{index}  {name}")
            TeachingPackageService._append_tracking_field(
                lines, "学习目标", item.get("learning_objective_ids")
            )
            TeachingPackageService._append_tracking_field(
                lines, "学习任务", item.get("student_task_ids")
            )
            TeachingPackageService._append_tracking_field(
                lines, "课堂评价", item.get("assessment_ids")
            )
            TeachingPackageService._append_tracking_field(
                lines, "教材证据", item.get("evidence_ids")
            )

    @staticmethod
    def _append_tracking_insufficiencies(lines: list[str], values: Any) -> None:
        if not values:
            return
        lines.extend(["", "证据与输入不足"])
        for index, item in enumerate(values, start=1):
            if index > 1:
                lines.append("")
            if not isinstance(item, dict):
                lines.append(f"不足{index}  {TeachingPackageService._display(item)}")
                continue
            name = TeachingPackageService._display(item.get("type") or "教材证据不足")
            lines.append(f"不足{index}  {name}")
            TeachingPackageService._append_tracking_field(
                lines, "说明", item.get("description")
            )
            TeachingPackageService._append_tracking_field(
                lines, "影响范围", item.get("affected_sections")
            )
            TeachingPackageService._append_tracking_field(
                lines, "建议操作", item.get("recommended_action")
            )

    @staticmethod
    def _append_tracking_field(lines: list[str], label: str, value: Any) -> None:
        if value not in ("", None, [], {}):
            lines.append(f"  {label}：{TeachingPackageService._display(value)}")

    def _load_context(self, request: TeachingPackageRequest) -> dict[str, Any]:
        with Session(self._engine) as session:
            course = session.get(CourseModel, request.course_id)
            if course is None:
                raise ValueError("课程不存在")
            document = session.get(DocumentModel, request.document_id)
            if document is None or document.course_id != request.course_id:
                raise ValueError("所选教材不属于当前课程")
            chapters = list(
                session.scalars(
                    select(ChapterModel)
                    .where(
                        ChapterModel.id.in_(request.chapter_ids),
                        ChapterModel.document_id == request.document_id,
                        ChapterModel.is_excluded.is_(False),
                    )
                    .order_by(ChapterModel.position)
                )
            )
            if {item.id for item in chapters} != set(request.chapter_ids):
                raise ValueError("所选章节不属于当前教材或已经失效")
            points = list(
                session.scalars(
                    select(KnowledgePointModel)
                    .where(
                        KnowledgePointModel.id.in_(request.knowledge_point_ids),
                        KnowledgePointModel.course_id == request.course_id,
                        KnowledgePointModel.status == "confirmed",
                        KnowledgePointModel.is_enabled.is_(True),
                    )
                    .order_by(KnowledgePointModel.id)
                )
            )
            if {item.id for item in points} != set(request.knowledge_point_ids):
                raise ValueError("部分知识点不属于当前课程、未确认或已停用")
            for item in (course, document, *chapters, *points):
                session.expunge(item)
        return {
            "course": course,
            "document": document,
            "chapters": chapters,
            "points": points,
        }

    def _collect_evidence(
        self,
        request: TeachingPackageRequest,
        points: list[KnowledgePointModel],
    ) -> tuple[TeachingEvidence, ...]:
        matches: dict[int, tuple[SearchResult, set[str]]] = {}
        for point in points:
            found = self._retriever.search(
                point.name,
                request.course_id,
                document_id=request.document_id,
                chapter_ids=list(request.chapter_ids),
                limit=4,
            )
            if not found and point.chapter_id in request.chapter_ids:
                found = self._retriever.scope_context(
                    request.course_id,
                    document_id=request.document_id,
                    chapter_ids=[point.chapter_id],
                    limit=2,
                )
            for item in found:
                if item.chunk_id in matches:
                    matches[item.chunk_id][1].add(point.name)
                else:
                    matches[item.chunk_id] = (item, {point.name})
        selected = list(matches.values())[:12]
        if not selected:
            return ()
        chunk_ids = [item.chunk_id for item, _ in selected]
        with Session(self._engine) as session:
            chunks = {
                chunk.id: chunk.content
                for chunk in session.scalars(
                    select(DocumentChunkModel).where(DocumentChunkModel.id.in_(chunk_ids))
                )
            }
        evidence: list[TeachingEvidence] = []
        for index, (item, point_names) in enumerate(selected, 1):
            content = chunks.get(item.chunk_id, item.excerpt).strip()
            evidence.append(
                TeachingEvidence(
                    evidence_id=f"E{index}",
                    chunk_id=item.chunk_id,
                    document_name=item.document_name,
                    chapter_title=item.chapter_title,
                    page_start=item.page_start,
                    page_end=item.page_end,
                    excerpt=content[:1800],
                    knowledge_points=tuple(sorted(point_names)),
                )
            )
        return tuple(evidence)

    @staticmethod
    def _build_user_prompt(
        request: TeachingPackageRequest,
        context: dict[str, Any],
        evidence: tuple[TeachingEvidence, ...],
    ) -> str:
        course: CourseModel = context["course"]
        document: DocumentModel = context["document"]
        chapters: list[ChapterModel] = context["chapters"]
        points: list[KnowledgePointModel] = context["points"]
        evidence_text = "\n\n".join(
            f"[{item.evidence_id}] 教材：{item.document_name}；"
            f"章节：{item.chapter_title}；页码：{item.page_start}-{item.page_end}；"
            f"关联知识点：{'、'.join(item.knowledge_points)}\n{item.excerpt}"
            for item in evidence
        )
        return (
            f"课程名称：{course.name}\n"
            f"学科：{course.subject}\n学段：{course.education_stage}\n"
            f"年级：{course.grade}\n教材名称：{document.filename}\n"
            f"教材版本：{course.textbook_version}\n"
            f"章节：{'、'.join(item.title for item in chapters)}\n"
            f"目标知识点：{'、'.join(item.name for item in points)}\n"
            f"课型：{request.lesson_type}\n"
            f"课时长度：{request.lesson_duration_minutes}分钟\n"
            f"学生情况：{request.student_profile or '未补充，请仅作一般性表述'}\n"
            f"教学侧重点：{request.teaching_focus or '按教材重点安排'}\n"
            f"教师补充要求：{request.additional_requirements or '无'}\n\n"
            f"教材证据：\n{evidence_text}"
        )

    @staticmethod
    def _validate_and_normalize(
        payload: dict[str, Any],
        request: TeachingPackageRequest,
        context: dict[str, Any],
        evidence: tuple[TeachingEvidence, ...],
    ) -> dict[str, Any]:
        if not isinstance(payload, dict):
            raise ValueError("模型返回的备课内容不是有效 JSON 对象")
        required_objects = (
            "material_tracking",
            "learning_guide",
            "teaching_plan",
            "quality_check",
        )
        missing = [name for name in required_objects if not isinstance(payload.get(name), dict)]
        if missing:
            raise ValueError("模型返回的备课内容缺少字段：" + "、".join(missing))
        payload["insufficiencies"] = (
            TeachingPackageService._normalize_insufficiencies(
                payload.get("insufficiencies")
            )
        )
        status = str(payload.get("status", "partial"))
        if status not in {"complete", "partial", "insufficient"}:
            status = "partial"
        valid_evidence_ids = {item.evidence_id for item in evidence}
        referenced_ids = TeachingPackageService._find_evidence_ids(payload)
        invented = referenced_ids - valid_evidence_ids
        if invented:
            raise ValueError("模型返回了不存在的教材证据编号：" + "、".join(sorted(invented)))

        course: CourseModel = context["course"]
        document: DocumentModel = context["document"]
        chapters: list[ChapterModel] = context["chapters"]
        title = str(payload.get("title") or f"{chapters[0].title}教学设计").strip()
        payload["status"] = status
        payload["title"] = title[:255]
        payload["basic_info"] = {
            "subject": course.subject,
            "school_stage": course.education_stage,
            "grade": course.grade,
            "textbook_name": document.filename,
            "textbook_version": course.textbook_version,
            "chapter_title": "、".join(item.title for item in chapters),
            "section_title": "、".join(item.title for item in chapters),
            "lesson_type": request.lesson_type,
            "lesson_duration_minutes": request.lesson_duration_minutes,
        }
        tracking = payload["material_tracking"]
        for key in ("knowledge_points", "coverage_matrix", "unsupported_knowledge_points"):
            if not isinstance(tracking.get(key), list):
                tracking[key] = []
        guide = payload["learning_guide"]
        guide.setdefault("document_name", "导学案")
        for key in (
            "learning_objectives",
            "key_points",
            "difficult_points",
            "prior_knowledge_check",
            "pre_class_preview",
            "knowledge_framework",
            "guided_fill_ins",
            "learning_tasks",
            "in_class_practice",
            "knowledge_summary",
            "after_class_tasks",
        ):
            if not isinstance(guide.get(key), list):
                guide[key] = []
        if not isinstance(guide.get("learning_summary"), dict):
            guide["learning_summary"] = {
                "knowledge_structure": [],
                "student_reflection_prompts": [],
            }
        plan = payload["teaching_plan"]
        plan.setdefault("document_name", "教案")
        plan.setdefault("textbook_analysis", "")
        plan.setdefault("student_analysis", "")
        for key in (
            "teaching_objectives",
            "key_points",
            "difficult_points",
            "teaching_methods",
            "learning_methods",
            "teaching_resources",
            "teaching_process",
            "board_design",
            "answer_reference",
            "homework_design",
            "post_lesson_reflection_template",
        ):
            if not isinstance(plan.get(key), list):
                plan[key] = []
        if not isinstance(plan.get("differentiated_instruction"), dict):
            plan["differentiated_instruction"] = {
                "support_for_struggling_students": [],
                "standard_requirements": [],
                "extension_for_advanced_students": [],
            }
        process = plan.get("teaching_process")
        assert isinstance(process, list)
        total_minutes = sum(
            TeachingPackageService._as_int(item.get("duration_minutes"))
            for item in process
            if isinstance(item, dict)
        )
        guide_task_ids = {
            str(item.get("id"))
            for key in ("pre_class_preview", "learning_tasks", "in_class_practice")
            for item in guide.get(key, [])
            if isinstance(item, dict) and item.get("id")
        }
        plan_task_ids = {
            str(task_id)
            for item in process
            if isinstance(item, dict)
            for task_id in (
                item.get("related_task_ids", [])
                if isinstance(item.get("related_task_ids"), list)
                else []
            )
        }
        plan_task_ids.update(
            identifier
            for item in process
            if isinstance(item, dict)
            for identifier in TeachingPackageService._identifier_tokens(
                item.get("assessment_id")
            )
        )
        unsupported = tracking.get("unsupported_knowledge_points")
        assert isinstance(unsupported, list)
        quality = payload["quality_check"]
        quality["all_knowledge_points_supported"] = not unsupported
        quality["guide_and_plan_aligned"] = guide_task_ids.issubset(plan_task_ids)
        quality["lesson_time_total_minutes"] = total_minutes
        quality["lesson_time_matches"] = total_minutes == request.lesson_duration_minutes
        guide_answer_items: list[dict[str, Any]] = []
        for key in (
            "prior_knowledge_check",
            "pre_class_preview",
            "knowledge_framework",
            "guided_fill_ins",
            "learning_tasks",
            "in_class_practice",
            "after_class_tasks",
        ):
            for item in guide.get(key, []):
                if not isinstance(item, dict):
                    continue
                has_question = bool(
                    item.get("question")
                    or item.get("questions")
                    or item.get("prompt")
                    or "______" in str(item.get("content") or "")
                    or "______" in str(item.get("instruction") or "")
                )
                if key in {
                    "guided_fill_ins",
                    "in_class_practice",
                    "after_class_tasks",
                } or has_question:
                    guide_answer_items.append(item)
        guide_answers_complete = all(
            str(item.get("answer") or "").strip()
            and str(item.get("analysis") or "").strip()
            for item in guide_answer_items
        )
        quality["answers_verified"] = bool(plan.get("answer_reference")) and bool(
            guide_answer_items
        ) and guide_answers_complete
        quality["out_of_scope_content_found"] = bool(
            quality.get("out_of_scope_content_found", False)
        )
        issues = quality.get("issues")
        if not isinstance(issues, list):
            issues = []
            quality["issues"] = issues
        if total_minutes != request.lesson_duration_minutes:
            issues.append(
                f"教学过程合计{total_minutes}分钟，与设定的"
                f"{request.lesson_duration_minutes}分钟不一致"
            )
            payload["status"] = "partial"
        if not quality["guide_and_plan_aligned"]:
            issues.append("导学案任务或评价编号未在教案教学过程中完整对应")
            payload["status"] = "partial"
        if guide_answer_items and not quality["answers_verified"]:
            issues.append("导学案练习、课后任务或教案中缺少答案与解析")
            payload["status"] = "partial"
        if len(guide.get("guided_fill_ins", [])) < 6:
            issues.append("导学案的知识梳理填空少于6项")
            payload["status"] = "partial"
        if not guide.get("knowledge_summary"):
            issues.append("导学案缺少知识点总结")
            payload["status"] = "partial"
        if unsupported and payload["status"] == "complete":
            payload["status"] = "partial"
        return payload

    @staticmethod
    def _as_int(value: Any) -> int:
        if isinstance(value, bool):
            return 0
        try:
            return int(value or 0)
        except (TypeError, ValueError):
            return 0

    @staticmethod
    def _normalize_insufficiencies(value: Any) -> list[dict[str, Any]]:
        if value is None or value == {}:
            return []
        if isinstance(value, str):
            text = value.strip()
            if not text or text.lower() in {"none", "null", "无", "暂无", "无不足"}:
                return []
            value = [{"description": text}]
        elif isinstance(value, dict):
            value = [value]
        elif not isinstance(value, list):
            return []

        normalized: list[dict[str, Any]] = []
        for item in value:
            if isinstance(item, str):
                item = {"description": item.strip()}
            if not isinstance(item, dict):
                continue
            description = str(item.get("description") or "").strip()
            item_type = str(item.get("type") or "教材或输入信息不足").strip()
            affected = item.get("affected_sections")
            if not isinstance(affected, list):
                affected = [str(affected)] if affected else []
            normalized.append(
                {
                    "type": item_type,
                    "description": description,
                    "affected_sections": affected,
                    "recommended_action": str(
                        item.get("recommended_action") or ""
                    ).strip(),
                }
            )
        return normalized

    @staticmethod
    def _identifier_tokens(value: Any) -> set[str]:
        if isinstance(value, list):
            return {
                token
                for item in value
                for token in TeachingPackageService._identifier_tokens(item)
            }
        if not isinstance(value, str):
            return set()
        return set(re.findall(r"[A-Z]+\d+", value.upper()))

    def _save(
        self,
        request: TeachingPackageRequest,
        payload: dict[str, Any],
        evidence: tuple[TeachingEvidence, ...],
    ) -> int:
        request_payload = asdict(request)
        request_payload["evidence"] = [asdict(item) for item in evidence]
        with Session(self._engine) as session:
            row = TeachingPackageModel(
                course_id=request.course_id,
                document_id=request.document_id,
                title=str(payload["title"]),
                status=str(payload["status"]),
                chapter_ids_json=json.dumps(request.chapter_ids),
                knowledge_point_ids_json=json.dumps(request.knowledge_point_ids),
                request_json=json.dumps(request_payload, ensure_ascii=False),
                result_json=json.dumps(payload, ensure_ascii=False),
                model_name=self._model_name,
            )
            session.add(row)
            session.commit()
            session.refresh(row)
            return row.id

    @staticmethod
    def _find_evidence_ids(value: Any) -> set[str]:
        found: set[str] = set()
        if isinstance(value, dict):
            for key, item in value.items():
                if key == "evidence_ids" and isinstance(item, list):
                    found.update(str(evidence_id) for evidence_id in item)
                else:
                    found.update(TeachingPackageService._find_evidence_ids(item))
        elif isinstance(value, list):
            for item in value:
                found.update(TeachingPackageService._find_evidence_ids(item))
        return found

    @staticmethod
    def _append_text_document(document: Any, content: str, title: str) -> None:
        section_titles = {
            "学习目标",
            "学习重点",
            "学习难点",
            "课前诊断",
            "课前预习",
            "知识框架",
            "知识梳理填空",
            "学习任务",
            "当堂练习",
            "知识点总结",
            "知识结构",
            "学习反思",
            "课后任务",
            "答案与解析",
            "课前诊断答案",
            "课前预习答案",
            "知识框架答案",
            "知识梳理填空答案",
            "学习任务答案",
            "当堂练习答案",
            "课后任务答案",
            "教材分析",
            "学情分析",
            "教学目标",
            "教学重点",
            "教学难点",
            "教学方法",
            "学习方法",
            "教学资源",
            "教学过程",
            "分层教学",
            "板书设计",
            "答案与评价参考",
            "作业设计",
            "课后反思",
            "知识点覆盖",
            "目标与任务对应",
            "教材原文依据",
            "证据与输入不足",
        }
        document.add_heading(title, level=1)
        lines = content.splitlines()
        if lines and lines[0].strip() == title:
            lines = lines[1:]
        for line in lines:
            text = line.strip()
            if not text:
                continue
            if text in section_titles:
                document.add_heading(text, level=2)
            else:
                document.add_paragraph(text)

    @staticmethod
    def _append_paragraph(lines: list[str], title: str, content: Any) -> None:
        if content:
            lines.extend(["", title, str(content)])

    @staticmethod
    def _append_simple_list(lines: list[str], title: str, values: Any) -> None:
        if not values:
            return
        lines.extend(["", title])
        if not isinstance(values, list):
            values = [values]
        for value in values:
            if isinstance(value, dict):
                text = TeachingPackageService._dict_summary(value)
            else:
                text = str(value)
            lines.append(f"• {text}")

    @staticmethod
    def _append_objectives(lines: list[str], title: str, values: Any) -> None:
        if not values:
            return
        lines.extend(["", title])
        for item in values:
            if not isinstance(item, dict):
                lines.append(f"• {item}")
                continue
            prefix = str(item.get("id") or "•")
            lines.append(f"{prefix}  {item.get('content', '')}")
            criteria = item.get("success_criteria")
            if criteria:
                lines.append(f"  达成标准：{criteria}")

    @staticmethod
    def _append_cards(
        lines: list[str],
        title: str,
        values: Any,
        hidden_keys: set[str] | None = None,
    ) -> None:
        if not values:
            return
        lines.extend(["", title])
        if not isinstance(values, list):
            values = [values]
        for index, item in enumerate(values):
            if index:
                lines.append("")
            if isinstance(item, dict):
                identifier = (
                    item.get("id")
                    or item.get("stage")
                    or item.get("title")
                    or item.get("task_or_assessment_id")
                    or "•"
                )
                primary = (
                    item.get("content")
                    or item.get("question")
                    or item.get("instruction")
                    or item.get("prompt")
                    or ""
                )
                rendered_keys: set[str] = set()
                if not primary:
                    answer = item.get("answer") or item.get("answers")
                    if answer not in ("", None, [], {}):
                        primary = f"答案：{TeachingPackageService._display(answer)}"
                        rendered_keys.update({"answer", "answers"})
                lines.append(f"{identifier}  {primary}".rstrip())
                for key, value in item.items():
                    if key in {
                        "id",
                        "stage",
                        "title",
                        "content",
                        "question",
                        "instruction",
                        "prompt",
                        "task_or_assessment_id",
                        "evidence_ids",
                        "knowledge_points",
                    } or key in rendered_keys or key in (hidden_keys or set()) or value in (
                        "",
                        None,
                        [],
                        {},
                    ):
                        continue
                    lines.append(
                        f"  {TeachingPackageService._label(key)}："
                        f"{TeachingPackageService._display(value)}"
                    )
            else:
                lines.append(f"• {item}")

    @staticmethod
    def _dict_summary(value: dict[str, Any]) -> str:
        return "；".join(
            f"{TeachingPackageService._label(key)}：{TeachingPackageService._display(item)}"
            for key, item in value.items()
            if item not in ("", None, [], {}) and key not in {"evidence_ids"}
        )

    @staticmethod
    def _display(value: Any) -> str:
        if isinstance(value, list):
            return "；".join(TeachingPackageService._display(item) for item in value)
        if isinstance(value, dict):
            return TeachingPackageService._dict_summary(value)
        return str(value)

    @staticmethod
    def _label(key: str) -> str:
        labels = {
            "purpose": "目的",
            "prompt": "填空",
            "question_type": "题型",
            "student_output": "学习成果",
            "questions": "问题",
            "learning_hint": "学习提示",
            "task_type": "任务类型",
            "scenario_or_material": "情境或材料",
            "difficulty": "难度",
            "score": "分值",
            "duration_minutes": "时间",
            "related_task_ids": "对应任务",
            "teacher_activities": "教师活动",
            "student_activities": "学生活动",
            "expected_student_responses": "预期表现",
            "possible_difficulties": "可能困难",
            "teacher_responses": "教学处理",
            "assessment_id": "评价编号",
            "assessment_method": "评价方式",
            "design_intention": "设计意图",
            "reason": "原因",
            "teaching_strategy": "教学策略",
            "breakthrough_strategy": "突破策略",
            "answer": "答案",
            "analysis": "简析",
            "solution_or_explanation": "解析",
            "scoring_criteria": "评价标准",
            "common_errors": "常见错误",
            "correction_strategy": "纠正策略",
            "level": "层级",
            "estimated_minutes": "预计时间",
            "knowledge_point": "知识点",
            "role": "作用",
            "teaching_level": "教学层级",
            "source_summary": "教材概述",
            "learning_objective_ids": "目标",
            "student_task_ids": "任务",
            "assessment_ids": "评价",
            "description": "说明",
            "affected_sections": "影响范围",
            "recommended_action": "建议操作",
        }
        return labels.get(key, key)
