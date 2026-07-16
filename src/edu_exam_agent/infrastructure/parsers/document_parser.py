"""Document text extraction for MVP file formats."""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol


@dataclass(frozen=True, slots=True)
class ParsedPage:
    page_number: int
    text: str


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    pages: tuple[ParsedPage, ...]
    page_errors: int = 0
    source_format: str = ""

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


class DocumentParser(Protocol):
    def parse(self, path: Path) -> ParsedDocument: ...


class PlainTextParser:
    def parse(self, path: Path) -> ParsedDocument:
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = path.read_text(encoding="gb18030")
        return ParsedDocument((ParsedPage(1, text),), source_format=path.suffix.lower())


class PdfParser:
    def parse(self, path: Path) -> ParsedDocument:
        try:
            return self._parse_with_pdfium(path)
        except Exception:
            return self._parse_with_pypdf(path)

    @staticmethod
    def _parse_with_pdfium(path: Path) -> ParsedDocument:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(path)
        pages: list[ParsedPage] = []
        page_errors = 0
        for index in range(len(document)):
            try:
                page = document[index]
                text_page = page.get_textpage()
                text = text_page.get_text_range()
                text_page.close()
                page.close()
            except Exception:
                text = ""
                page_errors += 1
            pages.append(ParsedPage(index + 1, text))
        document.close()
        return ParsedDocument(tuple(pages), page_errors=page_errors, source_format=".pdf")

    @staticmethod
    def _parse_with_pypdf(path: Path) -> ParsedDocument:
        from pypdf import PdfReader

        logger = logging.getLogger("pypdf")
        previous_level = logger.level
        logger.setLevel(logging.CRITICAL)
        try:
            reader = PdfReader(path, strict=False)
            pages: list[ParsedPage] = []
            page_errors = 0
            for index, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                    page_errors += 1
                pages.append(ParsedPage(index, text))
            return ParsedDocument(tuple(pages), page_errors=page_errors, source_format=".pdf")
        finally:
            logger.setLevel(previous_level)


class DocxParser:
    def parse(self, path: Path) -> ParsedDocument:
        from docx import Document

        document = Document(path)
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text.strip() for cell in row.cells))
        return ParsedDocument((ParsedPage(1, "\n".join(lines)),), source_format=".docx")


class ParserRegistry:
    """Resolve a parser without coupling workflows to a file library."""

    def __init__(self) -> None:
        plain = PlainTextParser()
        self._parsers: dict[str, DocumentParser] = {
            ".txt": plain,
            ".md": plain,
            ".markdown": plain,
            ".pdf": PdfParser(),
            ".docx": DocxParser(),
        }

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return tuple(self._parsers)

    def parse(self, path: Path) -> ParsedDocument:
        parser = self._parsers.get(path.suffix.lower())
        if parser is None:
            raise ValueError(f"不支持的教材格式：{path.suffix or '无扩展名'}")
        return parser.parse(path)
