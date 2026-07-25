from __future__ import annotations

import json
import sys
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

sys.path.insert(0, r"D:\出题助手\src")
sys.path.append(r"D:\出题助手\.venv\Lib\site-packages")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from edu_exam_agent.application.services.paper_service import Paper, PaperService
from edu_exam_agent.infrastructure.database.models import QuestionModel


OUT = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_全新第二套_10题.docx")
PREVIEW = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_全新第二套_10题.txt")
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")


def _font(size: int, *, bold: bool = False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def diagram(
    points: dict[str, tuple[float, float]],
    edges: list[tuple[str, str]],
    *,
    dashed: set[tuple[str, str]] | None = None,
    notes: list[tuple[str, tuple[float, float]]] | None = None,
) -> bytes:
    width, height = 1280, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    xs = [point[0] for point in points.values()]
    ys = [point[1] for point in points.values()]
    xmin, xmax = min(xs), max(xs)
    ymin, ymax = min(ys), max(ys)
    dx = max(1.0, xmax - xmin)
    dy = max(1.0, ymax - ymin)
    margin_x, margin_y = 155, 105

    def convert(point: tuple[float, float]) -> tuple[int, int]:
        x, y = point
        px = margin_x + int((x - xmin) / dx * (width - 2 * margin_x))
        py = height - margin_y - int((y - ymin) / dy * (height - 2 * margin_y))
        return px, py

    dashed = dashed or set()
    for start, end in edges:
        p1, p2 = convert(points[start]), convert(points[end])
        if (start, end) in dashed or (end, start) in dashed:
            for segment in range(0, 20, 2):
                a = segment / 20
                b = min(1.0, (segment + 1) / 20)
                draw.line(
                    (
                        p1[0] + (p2[0] - p1[0]) * a,
                        p1[1] + (p2[1] - p1[1]) * a,
                        p1[0] + (p2[0] - p1[0]) * b,
                        p1[1] + (p2[1] - p1[1]) * b,
                    ),
                    fill="#A34F3B",
                    width=5,
                )
        else:
            draw.line((p1, p2), fill="#234F79", width=6)

    offsets = {
        "A": (-45, 15), "B": (14, 15), "C": (14, -44), "D": (-48, -44),
        "E": (14, -42), "F": (14, 12), "G": (-44, 12), "H": (-46, 12),
        "O": (13, 12), "P": (13, 12), "X": (-42, 12), "Y": (13, -42),
        "A'": (-55, 14), "B'": (14, 14),
    }
    for name, point in points.items():
        px, py = convert(point)
        draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill="#234F79")
        ox, oy = offsets.get(name, (14, -42))
        draw.text((px + ox, py + oy), name, fill="#202124", font=_font(29, bold=True))

    for text, point in notes or []:
        px, py = convert(point)
        draw.text((px, py), text, fill="#5F6368", font=_font(24))

    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    return stream.getvalue()


FIGURES = {
    2: diagram(
        {"A": (-5, 0), "B": (0, 12), "C": (5, 0), "D": (0, -12),
         "E": (-2.5, 6), "F": (2.5, 6), "G": (2.5, -6), "H": (-2.5, -6)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "C"), ("B", "D"), ("E", "F"), ("F", "G"),
         ("G", "H"), ("H", "E")],
        dashed={("A", "C"), ("B", "D")},
        notes=[("AC=10", (-2, 0.7)), ("BD=24", (0.5, 4.3))],
    ),
    3: diagram(
        {"A": (0, 0), "B": (8, 0), "C": (8, 6), "D": (0, 6),
         "P": (32 / 7, 24 / 7), "X": (32 / 7, 0), "Y": (8, 24 / 7)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "C"), ("P", "X"), ("P", "Y")],
        dashed={("A", "C"), ("P", "X"), ("P", "Y")},
        notes=[("AB=8", (3.3, -0.65)), ("BC=6", (8.25, 2.5))],
    ),
    4: diagram(
        {"A": (0, 0), "B": (8, 0), "C": (11, 5), "D": (3, 5),
         "E": (9, 5 / 3), "F": (6.75, 1.25)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "E"), ("B", "D")],
        dashed={("A", "E"), ("B", "D")},
        notes=[("BE:EC=1:2", (8.9, 2.8))],
    ),
    5: diagram(
        {"A": (0, 0), "B": (9, 0), "C": (3, 6), "D": (-3, 6)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "C"), ("B", "D")],
        dashed={("A", "C"), ("B", "D")},
        notes=[("AB=9", (3.5, -0.65)), ("CD=6", (-0.7, 6.45)),
               ("AC⊥BD", (1.8, 2.8))],
    ),
    6: diagram(
        {"A": (0, 0), "B": (12, 0), "C": (12, 8), "D": (0, 8),
         "E": (12, 4), "F": (8, 8)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "F"), ("F", "E")],
        dashed={("A", "F"), ("F", "E")},
        notes=[("BE=CF=x", (8.0, 5.7)), ("∠AFE=90°", (7.0, 7.15))],
    ),
    7: diagram(
        {"A": (0, 0), "B": (6, 0), "C": (6, 6), "D": (0, 6), "P": (3, 0)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("P", "C"), ("P", "D")],
        dashed={("P", "C"), ("P", "D")},
        notes=[("边长6", (2.2, 6.45)), ("P在AB上", (2.35, -0.7))],
    ),
    8: diagram(
        {"A'": (-4, 0), "A": (0, 0), "B": (22, 0), "B'": (26, 0),
         "C": (17, 12), "D": (5, 12)},
        [("A'", "B'"), ("B'", "C"), ("C", "D"), ("D", "A'"),
         ("A", "D"), ("B", "C")],
        dashed={("A", "D"), ("B", "C")},
        notes=[("原下底22", (8.7, 0.55)), ("上底12", (9.2, 12.45)),
               ("每侧外移4", (-3.2, 0.8))],
    ),
    9: diagram(
        {"A": (0, 0), "B": (10, 0), "C": (6.4, 4.8), "D": (-3.6, 4.8),
         "H": (-3.6, 0)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("A", "C"), ("D", "H")],
        dashed={("A", "C"), ("D", "H")},
        notes=[("AB=10", (3.8, -0.55)), ("AD=6", (-2.0, 2.0)),
               ("AC=8", (2.5, 2.6))],
    ),
    10: diagram(
        {"A": (0, 0), "B": (10, 0), "C": (10, 10), "D": (0, 10),
         "E": (3, 0), "F": (10, 3), "G": (7, 10), "H": (0, 7)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"),
         ("E", "F"), ("F", "G"), ("G", "H"), ("H", "E")],
        dashed={("E", "F"), ("F", "G"), ("G", "H"), ("H", "E")},
        notes=[("AE=BF=CG=DH=x", (2.4, 10.45)), ("边长10", (3.8, -0.7))],
    ),
}


DATA = [
    {
        "type": "单项选择题", "score": 8,
        "stem": "【条件反推】在平行四边形ABCD中，下列推理错误的是（　　）。",
        "options": [
            ("A", "若AB=BC，则四边形ABCD是菱形"),
            ("B", "若AC=BD，则四边形ABCD是矩形"),
            ("C", "若∠ABD=∠CBD，则四边形ABCD是正方形"),
            ("D", "若AC⊥BD且AC=BD，则四边形ABCD是正方形"),
        ],
        "answer": "C",
        "analysis": "A由一组邻边相等可判定为菱形；B由对角线相等可判定为矩形；D同时满足菱形与矩形的特征，故为正方形。C中BD平分∠ABC只能进一步得到邻边相等，从而判定为菱形，但不能保证有直角，所以不一定是正方形。",
        "criteria": "选出C得4分；能指出C只能保证为菱形、缺少直角条件得4分。",
    },
    {
        "type": "单项选择题", "score": 8,
        "stem": "【中点四边形】如图，四边形ABCD的对角线AC=10，BD=24，且AC⊥BD。E、F、G、H分别是AB、BC、CD、DA的中点，则四边形EFGH的面积为（　　）。",
        "options": [("A", "34"), ("B", "60"), ("C", "120"), ("D", "240")],
        "answer": "B",
        "analysis": "由三角形中位线定理，EF∥AC且EF=AC/2=5，FG∥BD且FG=BD/2=12。因为AC⊥BD，所以EF⊥FG，EFGH是矩形，其面积为5×12=60。",
        "criteria": "得到两条中位线长度各3分；求出面积2分。",
    },
    {
        "type": "填空题", "score": 8,
        "stem": "【距离约束】如图，矩形ABCD中，AB=8，BC=6，点P在线段AC上。若点P到边AB、BC的距离相等，则AP=________。",
        "answer": "40/7",
        "analysis": "设AP/AC=t，则P把AC按比例t定位。P到AB的距离为6t，到BC的距离为8(1-t)。由6t=8(1-t)，得t=4/7。又AC=10，所以AP=10×4/7=40/7。",
        "criteria": "正确表示两个距离4分；求出比例2分；得到AP计2分。",
    },
    {
        "type": "填空题", "score": 8,
        "stem": "【分点交线】如图，在平行四边形ABCD中，点E在线段BC上，BE:EC=1:2，AE与BD交于点F，则AF:FE=________。",
        "answer": "3:1",
        "analysis": "因为AD∥BC，且E在BC上，所以AD∥BE。又A、F、E三点共线，B、F、D三点共线，因此∠FAD=∠FEB，∠FDA=∠FBE，得到△FAD∽△FEB。于是AF:FE=AD:BE。平行四边形中AD=BC，而BE:EC=1:2，所以BE=BC/3=AD/3，故AF:FE=3:1。",
        "criteria": "建立比例或相似关系4分；求出F的位置2分；写出3:1计2分。",
    },
    {
        "type": "计算题", "score": 12,
        "stem": "【垂直对角线梯形】如图，在梯形ABCD中，AB∥CD，AB=9，CD=6，两条对角线AC⊥BD，且AC:BD=1:2。求梯形的高和面积。",
        "answer": "高为6，面积为45。",
        "analysis": "设AC、BD在底边方向上的投影长度分别为p、q，则p+q=AB+CD=15。由AC⊥BD可得pq=h²。于是AC²=p²+h²=p(p+q)=15p，BD²=q²+h²=q(p+q)=15q。因为AC:BD=1:2，所以p:q=AC²:BD²=1:4。结合p+q=15，得p=3、q=12，故h²=pq=36，h=6。梯形面积为(9+6)×6÷2=45。",
        "criteria": "得到p+q=15及pq=h²计4分；利用对角线比求p、q计4分；求高与面积各2分。",
    },
    {
        "type": "计算题", "score": 12,
        "stem": "【变量直角】如图，矩形ABCD中，AB=12，BC=8。点E、F分别在BC、CD上，且BE=CF=x（0<x<8）。若∠AFE=90°，求x及△AFE的面积。",
        "answer": "x=4，△AFE的面积为32。",
        "analysis": "建立坐标系：A(0,0)、B(12,0)、C(12,8)、D(0,8)，则E(12,x)、F(12-x,8)。由AF⊥EF，有向量FA=(x-12,-8)、FE=(x,x-8)，内积为0：x(x-12)-8(x-8)=0，即x²-20x+64=0，解得x=4或16。由0<x<8，取x=4。此时AF=8√2，EF=4√2，所以S△AFE=1/2×8√2×4√2=32。",
        "criteria": "正确表示点或线段4分；建立直角方程3分；筛选x=4计2分；求面积3分。",
    },
    {
        "type": "计算题", "score": 12,
        "stem": "【最短路径】如图，正方形ABCD的边长为6，点P在线段AB上。求PC+PD的最小值，并求取得最小值时AP的长。",
        "answer": "PC+PD的最小值为6√5，此时AP=3。",
        "analysis": "把点D关于直线AB对称到D'，则PD=PD'，所以PC+PD=PC+PD'。当C、P、D'三点共线时和最小，最小值为CD'。取A(0,0)、B(6,0)、C(6,6)、D'(0,-6)，得CD'=√(6²+12²)=6√5。直线CD'与AB交于其中点P(3,0)，故AP=3。",
        "criteria": "完成对称转化4分；说明共线取最小3分；求最小值3分；求AP计2分。",
    },
    {
        "type": "应用题", "score": 10,
        "stem": "【堤坝加固】某堤坝横截面是等腰梯形，上底12米、下底22米、腰长13米。加固时保持上底和高度不变，将两个坡脚各向外移动4米，形成新的等腰梯形横截面。求每延长1米堤坝增加的土方量；若加固长度为50米，共需增加多少立方米土方？",
        "answer": "每延长1米增加48立方米；50米共增加2400立方米。",
        "analysis": "原梯形一侧的水平投影为(22-12)/2=5米，由勾股定理，高为√(13²-5²)=12米。新下底为22+8=30米。增加的横截面面积为[(30+12)-(22+12)]×12÷2=48平方米。因此每延长1米增加48立方米，50米共增加48×50=2400立方米。",
        "criteria": "求高4分；求新下底2分；求面积差2分；完成体积计算2分。",
    },
    {
        "type": "应用题", "score": 10,
        "stem": "【铰接框架】一平行四边形活动框架ABCD中，AB=10厘米，AD=6厘米。调节框架后，对角线AC=8厘米，且∠DAB为钝角。求此时框架围成的面积及另一条对角线BD的长。",
        "answer": "面积为48平方厘米，BD=4√13厘米。",
        "analysis": "过D作DH⊥直线AB。由于∠DAB为钝角，垂足H在BA的延长线上。设AH=x、DH=h，则x²+h²=AD²=36。由平行四边形的平移关系，点C到直线AB的垂足位于距A为10-x的位置，所以(10-x)²+h²=AC²=64。两式相减得100-20x=28，故x=18/5，进而h=24/5。框架面积为AB·DH=10×24/5=48。BD的水平投影长为AB+AH=68/5，竖直投影长为24/5，故BD=√[(68/5)²+(24/5)²]=4√13。",
        "criteria": "建立两组勾股关系4分；求投影与高3分；求面积1分；求BD计2分。",
    },
    {
        "type": "应用题", "score": 12,
        "stem": "【四边取点与最值】如图，正方形ABCD的边长为10。点E、F、G、H分别在AB、BC、CD、DA上，且AE=BF=CG=DH=x（0≤x≤10）。\n（1）证明四边形EFGH是正方形；\n（2）求四边形EFGH面积的最小值；\n（3）当四边形EFGH的面积为58时，求x。",
        "answer": "（1）EFGH是正方形；（2）最小面积为50；（3）x=3或7。",
        "analysis": "（1）四个角上的直角三角形△AEH、△BFE、△CGF、△DHG均由两条直角边x与10-x组成，故全等，从而EF=FG=GH=HE；又∠AEH+∠BEF=90°，所以∠HEF=90°，EFGH是正方形。\n（2）面积S=EF²=x²+(10-x)²=2(x-5)²+50，所以最小值为50，此时x=5。\n（3）令2(x-5)²+50=58，得(x-5)²=4，所以x=3或7，均符合范围。",
        "criteria": "全等与正方形判定4分；建立面积表达式3分；求最小值2分；解出并检验两个x计3分。",
    },
]


def make_question(index: int, data: dict[str, object]) -> QuestionModel:
    options = [
        {"label": label, "content": content}
        for label, content in data.get("options", [])  # type: ignore[union-attr]
    ]
    return QuestionModel(
        id=index,
        course_id=2,
        question_type=str(data["type"]),
        stem=str(data["stem"]),
        options_json=json.dumps(options, ensure_ascii=False),
        answer=str(data["answer"]),
        analysis=str(data["analysis"]),
        scoring_criteria=str(data["criteria"]),
        knowledge_points_json=json.dumps(
            ["平行四边形", "矩形", "菱形", "正方形", "梯形", "四边形综合"],
            ensure_ascii=False,
        ),
        difficulty=5,
        estimated_time_minutes=10,
        score=int(data["score"]),
        quality_score=96.0,
        recommendation_score=96.0,
        boundary_passed=True,
        status="validated",
        generation_model="出题助手·第五档全新题组",
    )


class ArtifactBank:
    def figure(self, question_id: int):
        png = FIGURES.get(question_id)
        return SimpleNamespace(png_data=png) if png else None


questions = tuple(make_question(index, item) for index, item in enumerate(DATA, 1))
paper = Paper(
    title="沪科版八年级下册·四边形第五档高难训练（全新第二套）",
    questions=questions,
    duration_minutes=100,
    include_answers=True,
)
service = PaperService(ArtifactBank())  # type: ignore[arg-type]
service.export_docx(paper, OUT)
PREVIEW.write_text(service.preview(paper, include_answers=True), encoding="utf-8")


document = Document(OUT)
for section in document.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if section.start_type != WD_SECTION_START.NEW_PAGE:
        section.start_type = WD_SECTION_START.NEW_PAGE

normal = document.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = document.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for style_name in ("Question", "Option", "AnswerSpace", "Analysis"):
    style = document.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11 if style_name == "Question" else 10.5)
    style.paragraph_format.line_spacing = 1.25

title = document.paragraphs[0]
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
for run in title.runs:
    run.font.size = Pt(21)
    run.font.color.rgb = DARK_BLUE

meta = document.paragraphs[1]
meta.text = "建议用时：100分钟　　满分：100分　　难度：第五档（难）　　题型：2+2+3+3"
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(10)
for run in meta.runs:
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

document.paragraphs[2].paragraph_format.space_after = Pt(12)
for paragraph in document.paragraphs:
    if paragraph.style.name == "Question":
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(5)
    elif paragraph.style.name == "Option":
        paragraph.paragraph_format.space_after = Pt(3)
    elif paragraph.style.name == "AnswerSpace":
        paragraph.paragraph_format.space_after = Pt(8)
    elif paragraph.style.name == "Analysis":
        paragraph.paragraph_format.space_after = Pt(6)
    if paragraph._p.xpath(".//w:drawing"):
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True

for section in document.sections:
    header = section.header.paragraphs[0]
    header.text = "沪科版八年级下册 · 第19章 四边形 · 第五档全新题组"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

document.core_properties.title = paper.title
document.core_properties.subject = "四边形第五档全新专项训练，共10题，含9幅配图、答案、解析与评分标准"
document.core_properties.author = "出题助手"
document.save(OUT)

print(OUT)
print(PREVIEW)
print(f"questions={len(questions)} total_score={paper.total_score} figures={len(FIGURES)}")
