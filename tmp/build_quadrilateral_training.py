from __future__ import annotations

import json
import sys
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

sys.path.insert(0, r"D:\出题助手\src")
sys.path.append(r"D:\出题助手\.venv\Lib\site-packages")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from edu_exam_agent.application.services.paper_service import Paper, PaperService
from edu_exam_agent.infrastructure.database.models import QuestionModel


OUT = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_10题.docx")
PREVIEW = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_10题.txt")
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


def _font(size: int, bold: bool = False):
    candidate = Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else FONT_PATH
    return ImageFont.truetype(str(candidate if candidate.exists() else FONT_PATH), size)


def diagram(
    points: dict[str, tuple[float, float]],
    edges: list[tuple[str, str]],
    *,
    dashed: set[tuple[str, str]] | None = None,
    notes: list[tuple[str, tuple[float, float]]] | None = None,
    axes: bool = False,
) -> bytes:
    width, height = 1280, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    xs = [value[0] for value in points.values()]
    ys = [value[1] for value in points.values()]
    xmin, xmax = min(xs), max(xs)
    ymin, ymax = min(ys), max(ys)
    dx = max(1.0, xmax - xmin)
    dy = max(1.0, ymax - ymin)
    margin_x, margin_y = 150, 105

    def convert(point: tuple[float, float]) -> tuple[int, int]:
        x, y = point
        px = margin_x + int((x - xmin) / dx * (width - margin_x * 2))
        py = height - margin_y - int((y - ymin) / dy * (height - margin_y * 2))
        return px, py

    if axes:
        origin = convert((0, 0))
        draw.line((55, origin[1], width - 55, origin[1]), fill="#687078", width=3)
        draw.line((origin[0], height - 45, origin[0], 45), fill="#687078", width=3)
        draw.text((width - 85, origin[1] + 18), "x", fill="#202124", font=_font(27))
        draw.text((origin[0] + 16, 42), "y", fill="#202124", font=_font(27))

    dashed = dashed or set()
    for start, end in edges:
        p1, p2 = convert(points[start]), convert(points[end])
        if (start, end) in dashed or (end, start) in dashed:
            segments = 18
            for index in range(0, segments, 2):
                a = index / segments
                b = min(1.0, (index + 1) / segments)
                draw.line(
                    (
                        p1[0] + (p2[0] - p1[0]) * a,
                        p1[1] + (p2[1] - p1[1]) * a,
                        p1[0] + (p2[0] - p1[0]) * b,
                        p1[1] + (p2[1] - p1[1]) * b,
                    ),
                    fill="#9A4E38",
                    width=5,
                )
        else:
            draw.line((p1, p2), fill="#234F79", width=6)

    label_font = _font(30, bold=True)
    offsets = {
        "A": (-42, 18), "B": (16, 18), "C": (16, -45), "D": (-48, -45),
        "E": (14, -42), "F": (-46, 18), "G": (14, 12), "H": (-46, -40),
        "O": (13, 13), "P": (-42, 20), "Q": (14, -42),
    }
    for name, value in points.items():
        px, py = convert(value)
        draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill="#234F79")
        ox, oy = offsets.get(name[0], (14, -42))
        draw.text((px + ox, py + oy), name, fill="#202124", font=label_font)

    note_font = _font(25)
    for text, position in notes or []:
        px, py = convert(position)
        draw.text((px, py), text, fill="#5F6368", font=note_font)

    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    return stream.getvalue()


FIGURES = {
    2: diagram(
        {"A": (0, 0), "B": (6, 0), "C": (11, 8.66), "D": (5, 8.66), "E": (9, 5.20), "H": (4.5, 7.79)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("A", "E"), ("E", "D"), ("E", "H")],
        dashed={("A", "E"), ("E", "H")},
        notes=[("AB=6", (2.4, -0.6)), ("AD=10", (1.8, 4.8)), ("60°", (0.75, 0.45))],
    ),
    3: diagram(
        {"A": (-7, 0), "B": (0, 5), "C": (7, 0), "D": (0, -5), "E": (-4, 0), "F": (4, 0), "O": (0, 0)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("A", "C"), ("B", "D"), ("B", "E"), ("E", "D"), ("D", "F"), ("F", "B")],
        dashed={("A", "C"), ("B", "D")},
        notes=[("AC=14", (-1.2, 0.65)), ("BD=10", (0.45, 2.0)), ("AE=CF=3", (-5.8, -0.9))],
    ),
    4: diagram(
        {"A": (0, 0), "B": (12, 0), "C": (12, 8), "D": (0, 8), "E": (3.06, 8)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("B", "E"), ("A", "E")],
        dashed={("B", "E"), ("A", "E")},
        notes=[("AB=12", (4.8, -0.7)), ("AD=8", (-0.9, 3.5)), ("A折到E", (4.3, 5.7))],
    ),
    5: diagram(
        {"A": (-8, 0), "B": (0, 6), "C": (8, 0), "D": (0, -6), "O": (0, 0), "E": (-4, 3)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("A", "C"), ("B", "D"), ("D", "E")],
        dashed={("A", "C"), ("B", "D"), ("D", "E")},
        notes=[("AC=16", (-1.5, 0.75)), ("BD=12", (0.55, 2.0)), ("E为AB中点", (-6.6, 3.8))],
    ),
    6: diagram(
        {"A": (0, 0), "B": (6, 0), "C": (6, 6), "D": (0, 6), "E": (6, 3), "F": (3, 0)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("A", "E"), ("D", "F"), ("E", "F")],
        dashed={("A", "E"), ("D", "F"), ("E", "F")},
        notes=[("DF⊥AE", (1.4, 4.2)), ("边长6", (2.3, 6.55))],
    ),
    7: diagram(
        {"A(-2,0)": (-2, 0), "B(2,4)": (2, 4), "C(6,0)": (6, 0), "D(2,m)": (2, -4)},
        [("A(-2,0)", "B(2,4)"), ("B(2,4)", "C(6,0)"), ("C(6,0)", "D(2,m)"), ("D(2,m)", "A(-2,0)")],
        axes=True,
    ),
    8: diagram(
        {"A": (0, 0), "B": (12, 0), "C": (12, 5), "D": (0, 5), "P": (8, 0), "Q": (4, 5)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("P", "C"), ("Q", "A")],
        dashed={("P", "C"), ("Q", "A")},
        notes=[("P：2单位/秒", (4.1, -0.75)), ("Q：1单位/秒", (2.2, 5.55))],
    ),
    9: diagram(
        {"A": (0, 0), "B": (9, 0), "C": (12, 6), "D": (3, 6), "E": (6.512, 4.025), "F": (7.146, 3.0), "G": (5.488, 1.975), "H": (4.854, 3.0)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("E", "F"), ("F", "G"), ("G", "H"), ("H", "E")],
        dashed={("E", "F"), ("F", "G"), ("G", "H"), ("H", "E")},
        notes=[("四个内角的角平分线", (3.0, 6.65))],
    ),
    10: diagram(
        {"A": (0, 0), "B": (8, 0), "C": (8, 6), "D": (0, 6), "E": (1.75, 0), "F": (6.25, 6)},
        [("A", "B"), ("B", "C"), ("C", "D"), ("D", "A"), ("D", "E"), ("E", "B"), ("B", "F"), ("F", "D")],
        dashed={("D", "E"), ("B", "F")},
        notes=[("AE=CF=x", (2.7, 6.55)), ("AB=8，BC=6", (2.4, -0.75))],
    ),
}


DATA = [
    {
        "type": "单项选择题", "score": 6,
        "stem": "【反例与判定】在四边形ABCD中，对角线AC、BD交于点O。下列条件中，能判定四边形ABCD一定是菱形的是（　　）。",
        "options": [
            ("A", "AC⊥BD，且AB=BC"),
            ("B", "OA=OC，且AC⊥BD"),
            ("C", "OA=OC，OB=OD，且AC⊥BD"),
            ("D", "AB=BC=CD，且∠B=∠D"),
        ],
        "answer": "C",
        "analysis": "C说明两条对角线互相平分，因此四边形ABCD先被判定为平行四边形；再由对角线互相垂直，可判定它是菱形。A、B均只给出一条对角线被平分，条件不足；D不能排除非平行四边形，可通过构造反例说明。",
        "criteria": "选出C得3分；能说明‘先判定平行四边形，再判定菱形’得3分。",
    },
    {
        "type": "计算题", "score": 8,
        "stem": "【角平分线构造】如图，在平行四边形ABCD中，AB=6，AD=10，∠DAB=60°。∠DAB的角平分线AE交BC于点E。求DE的长。",
        "answer": "DE=2√7。",
        "analysis": "因为AD∥BC，所以∠DAE=∠AEB；又AE平分∠DAB，所以∠DAE=∠EAB=30°，从而∠AEB=30°，AB=BE=6。于是EC=BC-BE=4。也可先在等腰三角形ABE中作高，求得AE=6√3。再从E向AD作垂线EH，在30°直角三角形AEH中，AH=9，EH=3√3，所以DH=AD-AH=1。由勾股定理，DE=√(DH²+EH²)=√28=2√7。",
        "criteria": "得到AB=BE及EC=4计3分；正确构造直角三角形计2分；求出DE计3分。",
    },
    {
        "type": "计算题", "score": 10,
        "stem": "【对角线中点模型】在平行四边形ABCD中，AC⊥BD，AC=14，BD=10。点E、F在线段AC上，且AE=CF=3。\n（1）证明四边形BEDF是平行四边形；\n（2）求四边形BEDF的周长。",
        "answer": "（1）四边形BEDF是平行四边形；（2）周长为4√41。",
        "analysis": "设AC、BD交于O。平行四边形的对角线互相平分，所以AO=CO=7，BO=DO=5。又AE=CF=3，故OE=OF=4。因此在四边形BEDF中，两条对角线BD、EF互相平分，四边形BEDF为平行四边形。又BD⊥EF，所以它还是菱形。在Rt△BOE中，BE=√(BO²+OE²)=√41，故周长为4√41。",
        "criteria": "找出O并说明两组中点关系4分；完成判定2分；勾股计算及周长4分。",
    },
    {
        "type": "应用题", "score": 8,
        "stem": "【矩形折叠】如图，矩形ABCD中，AB=12，AD=8。将矩形沿经过点B的一条直线折叠，使点A恰好落在边CD上的点E处。求DE的长。",
        "answer": "DE=12-4√5。",
        "analysis": "折叠保持对应线段长度，所以BE=BA=12。设DE=x，则CE=12-x。在Rt△BCE中，BC=8，由勾股定理得(12-x)²+8²=12²，所以12-x=4√5（E在线段CD上，取正值），因此x=12-4√5。",
        "criteria": "写出BE=BA计2分；建立勾股方程3分；正确取值并作答3分。",
    },
    {
        "type": "计算题", "score": 8,
        "stem": "【菱形与中点】如图，菱形ABCD的两条对角线AC=16，BD=12，点E是边AB的中点。求线段DE的长。",
        "answer": "DE=√97。",
        "analysis": "设两条对角线交于O。菱形的对角线互相垂直平分，取O为坐标原点、AC所在直线为x轴、BD所在直线为y轴，可设A(-8,0)、B(0,6)、D(0,-6)。因为E是AB中点，所以E(-4,3)。因此DE=√[(-4-0)²+(3+6)²]=√97。也可通过连接EO并运用中位线与勾股定理完成。",
        "criteria": "得到半对角线8、6计2分；正确表示中点位置3分；求出DE计3分。",
    },
    {
        "type": "计算题", "score": 10,
        "stem": "【正方形中的垂直转化】如图，正方形ABCD的边长为6，点E在边BC上。过点D作DF⊥AE，交AB于点F。\n（1）证明AF=BE；\n（2）若EF=3√2，求BE的长。",
        "answer": "（1）AF=BE；（2）BE=3。",
        "analysis": "（1）∠DAF=∠ABE=90°。又因为AD⊥AB、DF⊥AE，所以∠ADF=∠BAE；且AD=AB，故△ADF≌△BAE，得到AF=BE。\n（2）设AF=BE=x，则BF=6-x。由于BF⊥BE，Rt△FBE中EF²=BF²+BE²，即18=(6-x)²+x²，化简得(x-3)²=0，所以x=3。",
        "criteria": "找出两组等角3分；完成全等证明2分；建立方程3分；求解与检验2分。",
    },
    {
        "type": "计算题", "score": 10,
        "stem": "【坐标与判定】在平面直角坐标系中，A(-2,0)、B(2,4)、C(6,0)、D(2,m)。当四边形ABCD为平行四边形时，求m的值，并判断此时四边形ABCD是哪一种特殊平行四边形。",
        "answer": "m=-4，此时ABCD是正方形。",
        "analysis": "平行四边形的两条对角线互相平分。AC的中点为(2,0)，BD的中点为(2,(4+m)/2)，故(4+m)/2=0，得m=-4。此时AB²=(2+2)²+(4-0)²=32，BC²=(6-2)²+(0-4)²=32，且向量AB=(4,4)、BC=(4,-4)，内积为0，所以AB=BC且AB⊥BC。平行四边形ABCD为正方形。",
        "criteria": "利用中点关系求m计4分；说明邻边相等3分；说明有直角并判定3分。",
    },
    {
        "type": "应用题", "score": 10,
        "stem": "【动点与平行四边形】矩形ABCD中，AB=12，BC=5。点P从A出发，以每秒2个单位长度沿AB向B运动；点Q同时从D出发，以每秒1个单位长度沿DC向C运动。设运动时间为t秒（0≤t≤6）。\n（1）当t为何值时，四边形APCQ是平行四边形？\n（2）求此时四边形APCQ的周长。",
        "answer": "（1）t=4秒；（2）周长为16+2√41。",
        "analysis": "AP=2t，DQ=t，所以QC=12-t。AP∥QC；要使四边形APCQ为平行四边形，只需AP=QC，即2t=12-t，解得t=4，符合0≤t≤6。此时AP=8，PB=4。在Rt△PBC中，PC=√(PB²+BC²)=√41，所以周长为2(AP+PC)=16+2√41。",
        "criteria": "正确表示AP、QC计3分；列式求t计3分；勾股计算边长2分；周长2分。",
    },
    {
        "type": "计算题", "score": 10,
        "stem": "【角平分线综合】如图，在平行四边形ABCD中，∠A、∠B、∠C、∠D的角平分线依次相交于E、F、G、H。证明四边形EFGH是矩形。",
        "answer": "四边形EFGH是矩形。",
        "analysis": "平行四边形的对角相等，所以∠A与∠C的角平分线互相平行，∠B与∠D的角平分线互相平行。因此FG∥EH，EF∥GH，四边形EFGH是平行四边形。又平行四边形的邻角互补，即∠B+∠C=180°，所以它们的角平分线所成角为90°，即∠EFG=90°。有一个角是直角的平行四边形是矩形，故EFGH是矩形。",
        "criteria": "说明两组对边分别平行4分；证明相邻角平分线垂直4分；完成矩形判定2分。",
    },
    {
        "type": "应用题", "score": 12,
        "stem": "【参数与分类】矩形ABCD中，AB=8，BC=6。点E在边AB上，点F在边CD上，且AE=CF=x（0≤x<8）。\n（1）证明四边形DEBF总是平行四边形；\n（2）当x为何值时，四边形DEBF是菱形？\n（3）是否存在x，使四边形DEBF成为正方形？说明理由。",
        "answer": "（1）DEBF是平行四边形；（2）x=7/4；（3）不存在。",
        "analysis": "（1）EB=AB-AE=8-x，DF=DC-CF=8-x，所以EB=DF且EB∥DF，故DEBF是平行四边形。\n（2）在Rt△DAE中，DE=√(x²+36)。平行四边形DEBF为菱形时DE=EB，即√(x²+36)=8-x。两边非负，平方得x²+36=(8-x)²，解得x=7/4。\n（3）若DEBF为矩形，则DE⊥EB。EB与AB共线，而DE的水平分量为x，只有x=0时DE⊥EB；此时DE=6、EB=8，并非菱形，因此不是正方形。x=8时图形退化且不在范围内，所以不存在符合条件的x。",
        "criteria": "平行四边形判定3分；建立菱形方程并求解4分；分析直角条件3分；排除并作结论2分。",
    },
]


def make_question(index: int, data: dict[str, object]) -> QuestionModel:
    options = [
        {"label": label, "content": content}
        for label, content in data.get("options", [])  # type: ignore[union-attr]
    ]
    return QuestionModel(
        id=index,
        course_id=2,
        question_type=str(data["type"]),
        stem=str(data["stem"]),
        options_json=json.dumps(options, ensure_ascii=False),
        answer=str(data["answer"]),
        analysis=str(data["analysis"]),
        scoring_criteria=str(data["criteria"]),
        knowledge_points_json=json.dumps(
            ["平行四边形", "矩形", "菱形", "正方形", "四边形综合"],
            ensure_ascii=False,
        ),
        difficulty=5,
        estimated_time_minutes=9,
        score=int(data["score"]),
        quality_score=94.0,
        recommendation_score=94.0,
        boundary_passed=True,
        status="validated",
        generation_model="出题助手专项补题",
    )


class ArtifactBank:
    def figure(self, question_id: int):
        png = FIGURES.get(question_id)
        return SimpleNamespace(png_data=png) if png else None


type_order = {"单项选择题": 0, "填空题": 1, "计算题": 2, "应用题": 3}
questions = tuple(
    sorted(
        (make_question(index, item) for index, item in enumerate(DATA, 1)),
        key=lambda question: (type_order.get(question.question_type, 99), question.id),
    )
)
paper = Paper(
    title="沪科版八年级下册·四边形第五档高难训练",
    questions=questions,
    duration_minutes=90,
    include_answers=True,
)
service = PaperService(ArtifactBank())  # type: ignore[arg-type]
service.export_docx(paper, OUT)
PREVIEW.write_text(service.preview(paper, include_answers=True), encoding="utf-8")


document = Document(OUT)
for section in document.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if section.start_type != WD_SECTION_START.NEW_PAGE:
        section.start_type = WD_SECTION_START.NEW_PAGE

normal = document.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = document.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for style_name in ("Question", "Option", "AnswerSpace", "Analysis"):
    style = document.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11 if style_name == "Question" else 10.5)
    style.paragraph_format.line_spacing = 1.25

title = document.paragraphs[0]
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

meta = document.paragraphs[1]
meta.text = f"建议用时：90分钟　　满分：{paper.total_score}分　　难度：第五档（难）"
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(10)
for run in meta.runs:
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "Microsoft YaHei"
    )

identity = document.paragraphs[2]
identity.paragraph_format.space_after = Pt(12)

for paragraph in document.paragraphs:
    if paragraph.style.name == "Question":
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(5)
    elif paragraph.style.name == "Option":
        paragraph.paragraph_format.space_after = Pt(3)
    elif paragraph.style.name == "AnswerSpace":
        paragraph.paragraph_format.space_after = Pt(8)
    elif paragraph.style.name == "Analysis":
        paragraph.paragraph_format.space_after = Pt(6)
    if paragraph._p.xpath(".//w:drawing"):
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True

for section in document.sections:
    header = section.header.paragraphs[0]
    header.text = "沪科版八年级下册 · 第19章 四边形"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), "Microsoft YaHei"
        )

document.core_properties.title = paper.title
document.core_properties.subject = "八年级下册四边形第五档高难训练，共10题，含答案、解析与配图"
document.core_properties.author = "出题助手"
document.save(OUT)

print(OUT)
print(PREVIEW)
print(f"questions={len(questions)} total_score={paper.total_score} figures={len(FIGURES)}")
