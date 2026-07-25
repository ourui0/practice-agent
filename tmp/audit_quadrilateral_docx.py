from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

sys.path.append(r"D:\出题助手\.venv\Lib\site-packages")

from docx import Document
from docx.oxml.ns import qn


path = Path(r"D:\出题助手\output\沪科版八年级下册_四边形较难专项训练_10题.docx")
document = Document(path)
question_paragraphs = [
    paragraph for paragraph in document.paragraphs if paragraph.style.name == "Question"
]
question_section = question_paragraphs[:10]
answer_section = question_paragraphs[10:]

with zipfile.ZipFile(path) as archive:
    broken_member = archive.testzip()
    media = [name for name in archive.namelist() if name.startswith("word/media/")]

section_geometry = []
for section in document.sections:
    sect_pr = section._sectPr
    page_size = sect_pr.find(qn("w:pgSz"))
    margins = sect_pr.find(qn("w:pgMar"))
    section_geometry.append(
        {
            "page_width": page_size.get(qn("w:w")),
            "page_height": page_size.get(qn("w:h")),
            "top_margin": margins.get(qn("w:top")),
            "right_margin": margins.get(qn("w:right")),
            "bottom_margin": margins.get(qn("w:bottom")),
            "left_margin": margins.get(qn("w:left")),
        }
    )

assert broken_member is None
assert len(document.sections) == 2
assert len(document.inline_shapes) == 9
assert len(media) == 9
assert len(question_section) == 10
assert len(answer_section) == 10
assert all(f"{index}." in paragraph.text for index, paragraph in enumerate(question_section, 1))
assert all("答案：" in paragraph.text for paragraph in answer_section)
assert "参考答案与解析" in [paragraph.text for paragraph in document.paragraphs]
assert all(shape.width < document.sections[0].page_width for shape in document.inline_shapes)

print(
    json.dumps(
        {
            "file": str(path),
            "bytes": path.stat().st_size,
            "sections": len(document.sections),
            "question_paragraphs": len(question_section),
            "answer_paragraphs": len(answer_section),
            "figures": len(document.inline_shapes),
            "media_files": len(media),
            "zip_integrity": "ok",
            "section_geometry": section_geometry,
            "title": document.core_properties.title,
        },
        ensure_ascii=False,
        indent=2,
    )
)
