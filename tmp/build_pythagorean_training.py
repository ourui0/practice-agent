from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\出题助手\output\勾股定理提分训练_配图版V2.docx")
FIG_DIR = Path(r"D:\出题助手\tmp\pythagorean_figures")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


class _FakeAx:
    def __init__(self, fig):
        self.fig = fig
        self.commands = []
        self.xlim = (0, 1)
        self.ylim = (0, 1)

    def plot(self, xs, ys, style="-", lw=2, color="#234F79"):
        self.commands.append(("line", list(zip(xs, ys)), style, lw, color))

    def text(self, x, y, text, fontsize=10, color="#202124", weight=None):
        self.commands.append(("text", x, y, text, fontsize, color, weight))

    def scatter(self, xs, ys, s=35, color="#234F79"):
        colors = color if isinstance(color, list) else [color] * len(xs)
        for x, y, c in zip(xs, ys, colors):
            self.commands.append(("dot", x, y, max(4, int(s ** 0.5)), c))

    def axhline(self, y, color="#555555", lw=1.3):
        self.commands.append(("hline", y, lw, color))

    def axvline(self, x, color="#555555", lw=1.3):
        self.commands.append(("vline", x, lw, color))

    def annotate(self, _text, xy, xytext, arrowprops=None):
        color = (arrowprops or {}).get("color", "#234F79")
        lw = (arrowprops or {}).get("lw", 1.8)
        self.commands.append(("arrow", xytext, xy, lw, color))

    def set_xlim(self, a, b): self.xlim = (a, b)
    def set_ylim(self, a, b): self.ylim = (a, b)
    def set_aspect(self, *_args, **_kwargs): pass
    def axis(self, *_args, **_kwargs): pass


class _FakeFig:
    def __init__(self, figsize):
        self.figsize = figsize
        self.ax = _FakeAx(self)

    def tight_layout(self, **_kwargs): pass

    def savefig(self, path, dpi=220, **_kwargs):
        w, h = int(self.figsize[0] * dpi), int(self.figsize[1] * dpi)
        img = Image.new("RGB", (w, h), "white")
        draw = ImageDraw.Draw(img)
        ml, mr, mt, mb = 45, 45, 35, 40
        xmin, xmax = self.ax.xlim; ymin, ymax = self.ax.ylim
        def cv(p):
            return (ml + (p[0]-xmin)/(xmax-xmin)*(w-ml-mr), h-mb-(p[1]-ymin)/(ymax-ymin)*(h-mt-mb))
        font_path = r"C:\Windows\Fonts\msyh.ttc"
        for cmd in self.ax.commands:
            if cmd[0] == "hline":
                _, y, lw, color = cmd; p1,p2=cv((xmin,y)),cv((xmax,y)); draw.line([p1,p2],fill=color,width=max(1,int(lw*dpi/90)))
                continue
            if cmd[0] == "vline":
                _, x, lw, color = cmd; p1,p2=cv((x,ymin)),cv((x,ymax)); draw.line([p1,p2],fill=color,width=max(1,int(lw*dpi/90)))
                continue
            if cmd[0] == "line":
                _, pts, style, lw, color = cmd
                p1, p2 = cv(pts[0]), cv(pts[-1])
                width = max(1, int(lw*dpi/90))
                if style in ("--", ":"):
                    n = 18 if style == "--" else 30
                    for i in range(0, n, 2):
                        a=i/n; b=min((i+1)/n,1)
                        draw.line([(p1[0]+(p2[0]-p1[0])*a,p1[1]+(p2[1]-p1[1])*a),(p1[0]+(p2[0]-p1[0])*b,p1[1]+(p2[1]-p1[1])*b)],fill=color,width=width)
                else:
                    draw.line([p1,p2],fill=color,width=width)
            elif cmd[0] == "text":
                _, x, y, txt, size, color, weight = cmd
                px, py = cv((x,y)); font=ImageFont.truetype(font_path, max(14,int(size*dpi/55)))
                draw.text((px,py),txt,fill=color,font=font,anchor="ls",stroke_width=1 if weight=="bold" else 0)
            elif cmd[0] == "dot":
                _, x,y,r,color=cmd; px,py=cv((x,y)); rr=max(4,int(r*dpi/90)); draw.ellipse((px-rr,py-rr,px+rr,py+rr),fill=color)
            elif cmd[0] == "arrow":
                _, start,end,lw,color=cmd; p1,p2=cv(start),cv(end); width=max(2,int(lw*dpi/90)); draw.line([p1,p2],fill=color,width=width)
                import math
                ang=math.atan2(p2[1]-p1[1],p2[0]-p1[0]); L=16
                a=(p2[0]-L*math.cos(ang-0.45),p2[1]-L*math.sin(ang-0.45)); b=(p2[0]-L*math.cos(ang+0.45),p2[1]-L*math.sin(ang+0.45))
                draw.polygon([p2,a,b],fill=color)
        img.save(path)


class _FakePlt:
    def subplots(self, figsize=(4,3)):
        fig=_FakeFig(figsize); return fig, fig.ax
    def close(self, _fig): pass


plt = _FakePlt()


QUESTIONS = [
    ("1.【分类讨论·多选】", 6, "一个三角形的三边长分别为3，4，x，且它是直角三角形，则x的值可能是（　　）。\nA. √7　　B. 5　　C. √5　　D. 7"),
    ("2.【方程建模】", 8, "一个直角三角形的周长为30，斜边长为13。求这个直角三角形的两条直角边长。"),
    ("3.【作高转化】", 8, "在等腰三角形ABC中，AB＝AC＝13，BC＝10。点D是BC的中点。\n（1）求AD的长；\n（2）点E在线段AD上，且CE＝12，求AE的长。"),
    ("4.【折叠与方程】", 8, "在矩形ABCD中，AB＝8，BC＝6。将矩形沿经过点B的一条直线折叠，使点A恰好落在边CD上的点E处。求DE的长。"),
    ("5.【实际情境·双直角模型】", 8, "两根竖直旗杆的底端A、B在同一水平地面上，AB＝30米，两杆高分别为12米和28米。现从两杆顶端拉一条绷直的彩带，求彩带的长度。"),
    ("6.【空间展开·最短路径】", 10, "一个长方体盒子的长、宽、高分别为10 cm、8 cm、6 cm。一只蚂蚁从一个顶点出发，沿盒子表面爬到与它相对的顶点。求蚂蚁所走最短路线的长度。要求写出不同展开方式的比较过程。"),
    ("7.【坐标与距离】", 8, "在平面直角坐标系中，A（0，6），B（10，2）。点P在x轴上，且PA＝PB。求点P的坐标。"),
    ("8.【全等基础上的几何推理】", 10, "四边形ABCD中，∠BAD＝90°，AB＝AD＝5，BC＝5√2，CD＝10，且点A、C位于直线BD的两侧。\n（1）判断△BCD的形状，并说明理由；\n（2）求∠ABC的度数。"),
    ("9.【动点与勾股方程】", 10, "在矩形ABCD中，AB＝12，BC＝5。点P从A出发，以每秒2个单位长度的速度沿AB向B运动；同时点Q从C出发，以每秒1个单位长度的速度沿CB向B运动。设运动时间为t秒（0≤t≤5）。当PQ＝5时，求t。"),
    ("10.【探究题·整数直角三角形】", 12, "一个直角三角形的两条直角边和斜边都是正整数，且面积为30。求这个直角三角形的三边长，并说明为什么没有其他答案。"),
]


ANSWERS = [
    ("1", "A、B。", "若x为斜边，则x²＝3²＋4²＝25，得x＝5；若4为斜边，则3²＋x²＝4²，得x＝√7，且3＋√7＞4，能构成三角形。3不可能是斜边。", "每种分类2分，三边关系检验2分。"),
    ("2", "两条直角边分别为5和12。", "设两直角边为a、b。由周长得a＋b＝17；由勾股定理得a²＋b²＝169。于是(a＋b)²＝a²＋b²＋2ab，故289＝169＋2ab，得ab＝60。因此a、b是方程t²－17t＋60＝0的两根，解得t＝5或12。", "列出和与平方和关系3分，求出ab并建方程3分，答案2分。"),
    ("3", "（1）AD＝12；（2）AE＝12－√119。", "因为AB＝AC且D为BC中点，所以AD⊥BC，CD＝5。在Rt△ACD中，AD＝√(13²－5²)＝12。在Rt△CED中，DE＝√(12²－5²)＝√119。因为E在线段AD上，所以AE＝AD－DE＝12－√119。", "识别作高模型2分；两次正确使用勾股定理各3分。"),
    ("4", "DE＝8－2√7。", "折叠保持距离，因此BE＝BA＝8。设DE＝x，则CE＝8－x。在Rt△BCE中，BE²＝BC²＋CE²，所以64＝36＋(8－x)²。因E在线段CD上，8－x≥0，故8－x＝2√7，x＝8－2√7。", "写出BE＝BA 2分，列勾股方程3分，正确取值并作答3分。"),
    ("5", "34米。", "两杆顶端的竖直高度差为28－12＝16米，水平距离为30米。彩带、水平距离和高度差构成直角三角形，所以彩带长为√(30²＋16²)＝√1156＝34米。", "完成情境转化3分，列式3分，单位与答案2分。"),
    ("6", "2√74 cm。", "三种本质不同的展开得到的路线平方分别为：(6＋8)²＋10²＝296，(6＋10)²＋8²＝320，(8＋10)²＋6²＝360。因为296最小，所以最短路线长为√296＝2√74 cm。", "列出三种展开各2分，比较2分，化简与单位2分。"),
    ("7", "P（17/5，0）。", "设P（x，0）。由PA＝PB得PA²＝PB²，即x²＋6²＝(x－10)²＋2²。化简得20x＝68，所以x＝17/5。", "设点与距离平方关系3分，方程求解3分，坐标书写2分。"),
    ("8", "（1）△BCD是以∠DBC为直角的直角三角形；（2）∠ABC＝135°。", "在Rt△ABD中，BD²＝AB²＋AD²＝50，所以BD＝5√2。于是BD²＋BC²＝50＋50＝100＝CD²，由勾股定理的逆定理，△BCD为直角三角形，∠DBC＝90°。又AB＝AD且∠BAD＝90°，所以△ABD为等腰直角三角形，∠ABD＝45°。A、C在BD两侧，故∠ABC＝∠ABD＋∠DBC＝135°。", "求BD 2分，逆定理判断3分，求45° 2分，利用位置关系求角3分。"),
    ("9", "t＝18/5秒。", "运动t秒后，PB＝12－2t，BQ＝5－t，且∠PBQ＝90°。由PQ＝5得(12－2t)²＋(5－t)²＝25，化简为5t²－58t＋144＝0，解得t＝18/5或t＝8。因为0≤t≤5，所以舍去t＝8，故t＝18/5秒。", "表示线段2分，列方程3分，解方程3分，检验范围2分。"),
    ("10", "三边长为5、12、13。", "设两直角边为正整数a≤b。由面积为30，得ab＝60，因此只需检查因数对(1,60)、(2,30)、(3,20)、(4,15)、(5,12)、(6,10)。分别计算a²＋b²，只有5²＋12²＝169＝13²是完全平方数；其余各组的平方和都不是完全平方数。因此唯一答案是5、12、13。", "由面积得到ab＝60 2分，完整列出因数对4分，逐一判定4分，结论2分。"),
]


def finish_figure(fig, ax, name, xlim, ylim):
    ax.set_xlim(*xlim)
    ax.set_ylim(*ylim)
    ax.set_aspect("equal", adjustable="box")
    ax.axis("off")
    fig.tight_layout(pad=0.25)
    path = FIG_DIR / name
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def line(ax, p, q, style="-", width=2.0, color="#234F79"):
    ax.plot([p[0], q[0]], [p[1], q[1]], style, lw=width, color=color)


def label(ax, text, p, dx=0.12, dy=0.12, color="#202124"):
    ax.text(p[0] + dx, p[1] + dy, text, fontsize=11, color=color, weight="bold")


def make_figures():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    paths = {}

    # Q3: isosceles triangle with altitude and point E.
    fig, ax = plt.subplots(figsize=(4.2, 2.8))
    B, C, A, D, E = (0, 0), (10, 0), (5, 12), (5, 0), (5, 5.2)
    for p, q in ((A, B), (A, C), (B, C), (A, D), (C, E)):
        line(ax, p, q, "--" if (p, q) == (C, E) else "-")
    ax.plot([D[0], D[0] + 0.65, D[0] + 0.65], [D[1] + 0.65, D[1] + 0.65, D[1]], color="#234F79", lw=1.3)
    for t, p, dx, dy in (("A", A, 0.1, 0.1), ("B", B, -0.45, -0.65), ("C", C, 0.15, -0.65), ("D", D, -0.2, -0.7), ("E", E, -0.5, 0.05)):
        label(ax, t, p, dx, dy)
    ax.text(1.7, 6.2, "13", fontsize=10); ax.text(8.0, 6.2, "13", fontsize=10)
    ax.text(2.35, -0.85, "5", fontsize=10); ax.text(7.35, -0.85, "5", fontsize=10)
    ax.text(7.35, 2.7, "12", fontsize=10)
    paths[3] = finish_figure(fig, ax, "q03.png", (-1, 11), (-1.3, 13))

    # Q4: rectangle folding geometry.
    fig, ax = plt.subplots(figsize=(4.4, 3.0))
    A, B, C, D, E = (0, 0), (8, 0), (8, 6), (0, 6), (3.2, 6)
    for p, q in ((A, B), (B, C), (C, D), (D, A)):
        line(ax, p, q)
    line(ax, B, E, "--", 2.0, "#C05A3D")
    line(ax, A, E, ":", 1.6, "#7D8790")
    for t, p, dx, dy in (("A", A, -0.5, -0.55), ("B", B, 0.15, -0.55), ("C", C, 0.15, 0.1), ("D", D, -0.5, 0.1), ("E", E, -0.15, 0.18)):
        label(ax, t, p, dx, dy)
    ax.text(3.8, -0.7, "8", fontsize=10); ax.text(8.25, 2.8, "6", fontsize=10)
    ax.text(5.8, 3.2, "BE=BA", fontsize=10, color="#C05A3D")
    paths[4] = finish_figure(fig, ax, "q04.png", (-1, 10), (-1.2, 7.2))

    # Q5: two flagpoles and top cable.
    fig, ax = plt.subplots(figsize=(5.0, 2.8))
    A, B, U, V = (0, 0), (30, 0), (0, 12), (30, 28)
    line(ax, A, B, "-", 2.0, "#555555"); line(ax, A, U); line(ax, B, V)
    line(ax, U, V, "-", 2.3, "#C05A3D")
    line(ax, U, (30, 12), "--", 1.4, "#7D8790")
    for t, p, dx, dy in (("A", A, -1.2, -2), ("B", B, 0.5, -2), ("M", U, -1.5, 0.5), ("N", V, 0.5, 0.5)):
        label(ax, t, p, dx, dy)
    ax.text(13, -2.7, "30 m", fontsize=10); ax.text(-3.2, 5.5, "12 m", fontsize=10)
    ax.text(33, 18, "28 m", fontsize=10); ax.text(13, 13.3, "30 m", fontsize=10)
    ax.text(25.3, 20.5, "高度差16 m", fontsize=9, color="#7D8790")
    paths[5] = finish_figure(fig, ax, "q05.png", (-4, 36), (-3.5, 31))

    # Q6: cuboid and three surface unfold expressions.
    fig, ax = plt.subplots(figsize=(5.2, 3.1))
    A, B, C, D = (0, 0), (6, 0), (8.5, 2.2), (2.5, 2.2)
    A2, B2, C2, D2 = (0, 4), (6, 4), (8.5, 6.2), (2.5, 6.2)
    for p, q in ((A,B),(B,C),(C,D),(D,A),(A2,B2),(B2,C2),(C2,D2),(D2,A2),(A,A2),(B,B2),(C,C2),(D,D2)):
        line(ax,p,q,"--" if (p,q) in ((C,D),(D,A),(D,D2)) else "-")
    line(ax, A, C2, ":", 2.2, "#C05A3D")
    label(ax,"A",A,-0.55,-0.55); label(ax,"G",C2,0.15,0.15)
    ax.text(2.8,-0.65,"10",fontsize=10); ax.text(7.1,0.5,"8",fontsize=10); ax.text(6.35,2.7,"6",fontsize=10)
    ax.text(9.3,5.3,"展开后比较：",fontsize=10,weight="bold",color="#234F79")
    ax.text(9.3,4.1,"(6+8)²+10²",fontsize=10); ax.text(9.3,2.9,"(6+10)²+8²",fontsize=10); ax.text(9.3,1.7,"(8+10)²+6²",fontsize=10)
    paths[6] = finish_figure(fig, ax, "q06.png", (-1, 18), (-1.2, 7.2))

    # Q7: coordinate distance.
    fig, ax = plt.subplots(figsize=(5.0, 3.0))
    ax.axhline(0,color="#555555",lw=1.3); ax.axvline(0,color="#555555",lw=1.3)
    A, B, P = (0,6),(10,2),(3.4,0)
    line(ax,P,A,"--",1.8,"#234F79"); line(ax,P,B,"--",1.8,"#C05A3D")
    ax.scatter([A[0],B[0],P[0]],[A[1],B[1],P[1]],s=35,color=["#234F79","#C05A3D","#202124"])
    label(ax,"A(0,6)",A,0.2,0.2); label(ax,"B(10,2)",B,0.2,0.2); label(ax,"P(x,0)",P,-0.5,-0.9)
    ax.text(11.2,-0.4,"x",fontsize=10); ax.text(-0.45,7.4,"y",fontsize=10)
    paths[7] = finish_figure(fig, ax, "q07.png", (-1.3, 13), (-1.3, 8))

    # Q8: quadrilateral split by BD.
    fig, ax = plt.subplots(figsize=(4.2, 3.0))
    B, D, A, C = (0,0),(6,0),(3,-3),(0,6)
    for p,q in ((A,B),(A,D),(B,C),(C,D),(B,D)):
        line(ax,p,q,"--" if (p,q)==(B,D) else "-")
    for t,p,dx,dy in (("A",A,-0.15,-0.8),("B",B,-0.6,-0.3),("C",C,-0.5,0.2),("D",D,0.2,-0.3)):
        label(ax,t,p,dx,dy)
    ax.text(1.1,-1.7,"5",fontsize=10); ax.text(4.6,-1.7,"5",fontsize=10)
    ax.text(-1.05,2.7,"5√2",fontsize=10); ax.text(3.5,3.15,"10",fontsize=10)
    paths[8] = finish_figure(fig, ax, "q08.png", (-1.5, 7.5), (-4, 7))

    # Q9: moving points in a rectangle.
    fig, ax = plt.subplots(figsize=(5.0, 2.7))
    A,B,C,D,P,Q=(0,0),(12,0),(12,5),(0,5),(5.2,0),(12,2.6)
    for p,q in ((A,B),(B,C),(C,D),(D,A)):
        line(ax,p,q)
    line(ax,P,Q,"--",2.1,"#C05A3D")
    ax.annotate("",xy=(8.2,0),xytext=(4.8,0),arrowprops=dict(arrowstyle="->",color="#234F79",lw=1.8))
    ax.annotate("",xy=(12,1.4),xytext=(12,3.7),arrowprops=dict(arrowstyle="->",color="#234F79",lw=1.8))
    for t,p,dx,dy in (("A",A,-0.55,-0.55),("B",B,0.15,-0.55),("C",C,0.15,0.15),("D",D,-0.55,0.15),("P",P,-0.15,-0.65),("Q",Q,0.2,0)):
        label(ax,t,p,dx,dy)
    ax.text(5.6,-1.0,"12",fontsize=10); ax.text(12.35,4.0,"5",fontsize=10)
    ax.text(6.4,0.45,"2单位/秒",fontsize=9,color="#234F79"); ax.text(9.0,3.1,"1单位/秒",fontsize=9,color="#234F79")
    paths[9] = finish_figure(fig, ax, "q09.png", (-1, 15), (-1.3, 6.3))
    return paths


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)


FIGURES = make_figures()
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, attr, Inches(1))
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
):
    st = doc.styles[style_name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("沪科版八年级下册 · 第18章"), 9, color=GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("勾股定理提分训练"), 9, color=GRAY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("勾股定理提分训练"), 22, True, DARK)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(14)
set_font(sub.add_run("综合提升 · 构造、方程、分类与最短路径"), 12, False, GRAY)

meta = doc.add_table(rows=2, cols=4)
meta.autofit = False
widths = [Inches(0.75), Inches(2.45), Inches(0.75), Inches(2.55)]
for row in meta.rows:
    for cell, width in zip(row.cells, widths):
        cell.width = width
meta.cell(0, 0).text = "姓名"
meta.cell(0, 1).text = ""
meta.cell(0, 2).text = "用时"
meta.cell(0, 3).text = "建议55分钟"
meta.cell(1, 0).text = "得分"
meta.cell(1, 1).text = ""
meta.cell(1, 2).text = "满分"
meta.cell(1, 3).text = "88分"
for row in meta.rows:
    for i, cell in enumerate(row.cells):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_font(run, 10.5, bold=i in (0, 2))

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
note.paragraph_format.space_after = Pt(10)
set_font(note.add_run("作答提示："), 10.5, True, DARK)
set_font(note.add_run("先判断直角三角形在哪里，再决定使用勾股定理还是逆定理；含参数、动点问题注意分类或取值范围。"), 10.5)

doc.add_heading("训练题", level=1)
for question_no, (title_text, score, stem) in enumerate(QUESTIONS, 1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title_text), 11, True, DARK)
    set_font(p.add_run(f"（{score}分）"), 10.5, False, GRAY)
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(4)
    set_font(body.add_run(stem), 11)
    if question_no in FIGURES:
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(2)
        pic.paragraph_format.space_after = Pt(4)
        pic.add_run().add_picture(str(FIGURES[question_no]), width=Inches(4.7))
    lines = 1 if score <= 6 else (2 if score <= 8 else 3)
    for _ in range(lines):
        ans = doc.add_paragraph()
        ans.paragraph_format.space_after = Pt(3)
        set_font(ans.add_run("答：" if _ == 0 else "　"), 10, color=GRAY)

doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading("参考答案与解析", level=1)
intro = doc.add_paragraph("建议先独立完成，再按“模型识别—关系建立—运算—检验”四个环节订正。")
for run in intro.runs:
    set_font(run, 10.5, color=GRAY)

for number, answer, analysis, scoring in ANSWERS:
    h = doc.add_heading(f"第{number}题", level=2)
    h.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("答案："), 11, True, DARK)
    set_font(p.add_run(answer), 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("解析："), 11, True, DARK)
    set_font(p.add_run(analysis), 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("评分建议："), 10.5, True, GRAY)
    set_font(p.add_run(scoring), 10.5, color=GRAY)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
