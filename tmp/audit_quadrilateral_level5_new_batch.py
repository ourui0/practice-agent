from __future__ import annotations

import re
import sys
import zipfile
from difflib import SequenceMatcher
from pathlib import Path

sys.path.append(r"D:\出题助手\.venv\Lib\site-packages")

from docx import Document


output = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_全新第二套_10题.docx")
old_output = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_10题.docx")
document = Document(output)
texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

headings = [text for text in texts if text.startswith(("一、", "二、", "三、", "四、"))]
answers = [text for text in texts if re.match(r"^\d+\. 答案：", text)]
analyses = [text for text in texts if text.startswith("解析：")]
criteria = [text for text in texts if text.startswith("评分标准：")]
with zipfile.ZipFile(output) as archive:
    media = [name for name in archive.namelist() if name.startswith("word/media/")]

assert headings[:4] == [
    "一、选择题（共2题）",
    "二、填空题（共2题）",
    "三、计算题（共3题）",
    "四、应用题（共3题）",
]
assert len(answers) == len(analyses) == len(criteria) == 10
assert len(media) == 9
assert any("难度：第五档（难）" in text for text in texts)
assert any(text.startswith("10. 答案：") for text in texts)
assert not any("表达需统一" in text or "TODO" in text for text in texts)


def stems(path: Path) -> list[str]:
    return [
        paragraph.text
        for paragraph in Document(path).paragraphs
        if paragraph.style.name == "Question" and "答案：" not in paragraph.text
    ]


new_stems = stems(output)
old_stems = stems(old_output)
assert len(new_stems) == 10
assert len(old_stems) == 10
comparisons = [
    (SequenceMatcher(None, new, old).ratio(), new, old)
    for new in new_stems
    for old in old_stems
]
max_ratio, closest_new, closest_old = max(comparisons, key=lambda item: item[0])

print(f"headings={headings[:4]}")
print(f"answers={len(answers)} analyses={len(analyses)} criteria={len(criteria)} media={len(media)}")
print(f"sections={len(document.sections)} paragraphs={len(document.paragraphs)} size={output.stat().st_size}")
print(f"max_old_new_similarity={max_ratio:.3f}")
print(f"closest_new={closest_new}")
print(f"closest_old={closest_old}")
assert max_ratio < 0.70
