from __future__ import annotations

import sys
import zipfile
from pathlib import Path

sys.path.append(r"D:\出题助手\.venv\Lib\site-packages")

from docx import Document


path = Path(r"D:\出题助手\output\沪科版八年级下册_四边形第五档高难训练_10题.docx")
document = Document(path)
texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
headings = [text for text in texts if text.startswith(("一、", "二、", "三、", "四、"))]
answers = [text for text in texts if "答案：" in text]
analyses = [text for text in texts if text.startswith("解析：")]
criteria = [text for text in texts if text.startswith("评分标准：")]
with zipfile.ZipFile(path) as archive:
    media = [name for name in archive.namelist() if name.startswith("word/media/")]

assert headings[:3] == [
    "一、选择题（共1题）",
    "二、计算题（共6题）",
    "三、应用题（共3题）",
]
assert len(answers) == 10
assert len(media) == 9
assert any("难度：第五档（难）" in text for text in texts)
assert any(text.startswith("10. 答案：") for text in texts)
assert len(analyses) == 10
assert len(criteria) == 10

print(f"headings={headings[:3]}")
print(
    f"answers={len(answers)} analyses={len(analyses)} "
    f"criteria={len(criteria)} media={len(media)}"
)
print(f"sections={len(document.sections)} paragraphs={len(document.paragraphs)}")
print(f"size={path.stat().st_size}")
